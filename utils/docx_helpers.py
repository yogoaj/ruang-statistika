"""
utils/docx_helpers.py — Word document generation, multi-style
Ruang Statistika v6.1

Perbaikan v6.1:
- TIDAK ada _add_blank() — jarak antar elemen pakai space_before/space_after saja
- Heading pakai space_before=Pt(18) dan space_after=Pt(6), bukan blank paragraf
- Spacer tabel hanya space_after=Pt(6) pada elemen terakhir tabel
- Narasi AI fallback: jika ai_texts[modul] kosong, dibuat narasi otomatis dari data statistik
- Paragraf kosong hanya muncul di title page (by design)

Style didukung:
  APA 7th Edition | Skripsi/Tesis Indonesia (DIKTI) | Vancouver (Medis/Kesehatan)
  Jurnal Ilmiah Umum | Laporan Bisnis/Kantor
"""

import io
import datetime
import re

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _first_valid_df(*candidates):
    """Kembalikan DataFrame/nilai pertama yang tidak None dan tidak kosong.
    Aman digunakan sebagai pengganti `a or b` ketika salah satu bisa berupa DataFrame."""
    for c in candidates:
        if c is None:
            continue
        try:
            if not c.empty:
                return c
        except AttributeError:
            if c:
                return c
    return None



# =============================================================================
# STYLE PROFILES
# =============================================================================

class StyleProfile:
    def __init__(
        self,
        font, font_size_pt, line_spacing_mult,
        margin_cm, left_margin_cm, first_line_indent_cm,
        para_space_after_pt,          # jarak antar paragraf (pengganti blank line)
        h1_center, h1_bold, h1_upper, h1_color,
        h1_space_before_pt, h1_space_after_pt,
        h2_bold, h2_italic, h2_color,
        h2_space_before_pt, h2_space_after_pt,
        tbl_has_color_header, tbl_header_hex, tbl_header_txt_hex,
        tbl_font_size_pt, tbl_space_before_pt, tbl_space_after_pt,
        tbl_caption_above,
        tbl_caption_bold_label, tbl_caption_italic_desc,
        fig_space_after_pt,
        fig_caption_bold_label, fig_caption_italic_desc,
        ref_hanging_indent, ref_label,
        ai_show_label, ai_label_level,
    ):
        self.font                   = font
        self.font_size              = Pt(font_size_pt)
        self.line_spacing           = line_spacing_mult
        self.margin                 = Cm(margin_cm)
        self.left_margin            = Cm(left_margin_cm)
        self.first_line_indent      = Cm(first_line_indent_cm)
        self.para_space_after       = Pt(para_space_after_pt)
        self.h1_center              = h1_center
        self.h1_bold                = h1_bold
        self.h1_upper               = h1_upper
        self.h1_color               = RGBColor(*h1_color)
        self.h1_space_before        = Pt(h1_space_before_pt)
        self.h1_space_after         = Pt(h1_space_after_pt)
        self.h2_bold                = h2_bold
        self.h2_italic              = h2_italic
        self.h2_color               = RGBColor(*h2_color)
        self.h2_space_before        = Pt(h2_space_before_pt)
        self.h2_space_after         = Pt(h2_space_after_pt)
        self.tbl_has_color_header   = tbl_has_color_header
        self.tbl_header_hex         = tbl_header_hex
        self.tbl_header_txt_hex     = tbl_header_txt_hex
        self.tbl_font_size          = Pt(tbl_font_size_pt)
        self.tbl_space_before       = Pt(tbl_space_before_pt)
        self.tbl_space_after        = Pt(tbl_space_after_pt)
        self.tbl_caption_above      = tbl_caption_above
        self.tbl_caption_bold_label = tbl_caption_bold_label
        self.tbl_caption_italic_desc= tbl_caption_italic_desc
        self.fig_space_after        = Pt(fig_space_after_pt)
        self.fig_caption_bold_label = fig_caption_bold_label
        self.fig_caption_italic_desc= fig_caption_italic_desc
        self.ref_hanging_indent     = ref_hanging_indent
        self.ref_label              = ref_label
        self.ai_show_label          = ai_show_label
        self.ai_label_level         = ai_label_level


STYLE_PROFILES = {
    # ------------------------------------------------------------------
    # APA 7th Edition
    # ------------------------------------------------------------------
    "APA 7th Edition": StyleProfile(
        font="Times New Roman", font_size_pt=12,
        line_spacing_mult=2.0,
        margin_cm=2.54, left_margin_cm=2.54,
        first_line_indent_cm=1.27,
        para_space_after_pt=0,
        h1_center=True, h1_bold=True, h1_upper=False, h1_color=(0,0,0),
        h1_space_before_pt=24, h1_space_after_pt=0,
        h2_bold=True, h2_italic=False, h2_color=(0,0,0),
        h2_space_before_pt=12, h2_space_after_pt=0,
        tbl_has_color_header=False,
        tbl_header_hex="FFFFFF", tbl_header_txt_hex="000000",
        tbl_font_size_pt=12,
        tbl_space_before_pt=12, tbl_space_after_pt=12,
        tbl_caption_above=True,
        tbl_caption_bold_label=True, tbl_caption_italic_desc=True,
        fig_space_after_pt=12,
        fig_caption_bold_label=True, fig_caption_italic_desc=True,
        ref_hanging_indent=True, ref_label="References",
        ai_show_label=True, ai_label_level=3,
    ),

    # ------------------------------------------------------------------
    # Skripsi / Tesis Indonesia (DIKTI)
    # ------------------------------------------------------------------
    "Skripsi / Tesis Indonesia (DIKTI)": StyleProfile(
        font="Times New Roman", font_size_pt=12,
        line_spacing_mult=2.0,
        margin_cm=3.0, left_margin_cm=4.0,
        first_line_indent_cm=1.25,
        para_space_after_pt=0,
        h1_center=True, h1_bold=True, h1_upper=True, h1_color=(0,0,0),
        h1_space_before_pt=24, h1_space_after_pt=0,
        h2_bold=True, h2_italic=False, h2_color=(0,0,0),
        h2_space_before_pt=12, h2_space_after_pt=0,
        tbl_has_color_header=False,
        tbl_header_hex="FFFFFF", tbl_header_txt_hex="000000",
        tbl_font_size_pt=11,
        tbl_space_before_pt=12, tbl_space_after_pt=12,
        tbl_caption_above=True,
        tbl_caption_bold_label=True, tbl_caption_italic_desc=False,
        fig_space_after_pt=12,
        fig_caption_bold_label=True, fig_caption_italic_desc=False,
        ref_hanging_indent=True, ref_label="Daftar Pustaka",
        ai_show_label=True, ai_label_level=3,
    ),

    # ------------------------------------------------------------------
    # Vancouver (Medis / Kesehatan)
    # ------------------------------------------------------------------
    "Vancouver (Medis / Kesehatan)": StyleProfile(
        font="Arial", font_size_pt=11,
        line_spacing_mult=1.5,
        margin_cm=2.5, left_margin_cm=2.5,
        first_line_indent_cm=0.0,
        para_space_after_pt=6,
        h1_center=False, h1_bold=True, h1_upper=False, h1_color=(44,110,73),
        h1_space_before_pt=18, h1_space_after_pt=6,
        h2_bold=True, h2_italic=False, h2_color=(44,110,73),
        h2_space_before_pt=12, h2_space_after_pt=4,
        tbl_has_color_header=True,
        tbl_header_hex="2C6E49", tbl_header_txt_hex="FFFFFF",
        tbl_font_size_pt=10,
        tbl_space_before_pt=10, tbl_space_after_pt=10,
        tbl_caption_above=True,
        tbl_caption_bold_label=True, tbl_caption_italic_desc=True,
        fig_space_after_pt=10,
        fig_caption_bold_label=False, fig_caption_italic_desc=True,
        ref_hanging_indent=False, ref_label="References",
        ai_show_label=True, ai_label_level=3,
    ),

    # ------------------------------------------------------------------
    # Jurnal Ilmiah Umum
    # ------------------------------------------------------------------
    "Jurnal Ilmiah Umum": StyleProfile(
        font="Times New Roman", font_size_pt=11,
        line_spacing_mult=1.1,
        margin_cm=2.5, left_margin_cm=2.5,
        first_line_indent_cm=0.5,
        para_space_after_pt=4,
        h1_center=False, h1_bold=True, h1_upper=True, h1_color=(74,74,74),
        h1_space_before_pt=16, h1_space_after_pt=4,
        h2_bold=True, h2_italic=False, h2_color=(74,74,74),
        h2_space_before_pt=10, h2_space_after_pt=2,
        tbl_has_color_header=True,
        tbl_header_hex="4A4A4A", tbl_header_txt_hex="FFFFFF",
        tbl_font_size_pt=10,
        tbl_space_before_pt=8, tbl_space_after_pt=8,
        tbl_caption_above=True,
        tbl_caption_bold_label=True, tbl_caption_italic_desc=True,
        fig_space_after_pt=8,
        fig_caption_bold_label=True, fig_caption_italic_desc=True,
        ref_hanging_indent=True, ref_label="References",
        ai_show_label=True, ai_label_level=3,
    ),

    # ------------------------------------------------------------------
    # Laporan Bisnis / Kantor
    # ------------------------------------------------------------------
    "Laporan Bisnis / Kantor": StyleProfile(
        font="Calibri", font_size_pt=11,
        line_spacing_mult=1.15,
        margin_cm=2.5, left_margin_cm=2.5,
        first_line_indent_cm=0.0,
        para_space_after_pt=8,
        h1_center=False, h1_bold=True, h1_upper=False, h1_color=(24,95,165),
        h1_space_before_pt=16, h1_space_after_pt=6,
        h2_bold=True, h2_italic=False, h2_color=(24,95,165),
        h2_space_before_pt=10, h2_space_after_pt=4,
        tbl_has_color_header=True,
        tbl_header_hex="185FA5", tbl_header_txt_hex="FFFFFF",
        tbl_font_size_pt=10,
        tbl_space_before_pt=8, tbl_space_after_pt=8,
        tbl_caption_above=False,
        tbl_caption_bold_label=True, tbl_caption_italic_desc=False,
        fig_space_after_pt=10,
        fig_caption_bold_label=True, fig_caption_italic_desc=False,
        ref_hanging_indent=False, ref_label="Referensi",
        ai_show_label=False, ai_label_level=3,
    ),
}

_DEFAULT_PROFILE = STYLE_PROFILES["APA 7th Edition"]


def _get_profile(report_style: str) -> StyleProfile:
    return STYLE_PROFILES.get(report_style, _DEFAULT_PROFILE)


# =============================================================================
# DOCUMENT SETUP
# =============================================================================

def _set_doc_defaults(doc: Document, p: StyleProfile):
    for section in doc.sections:
        section.top_margin    = p.margin
        section.bottom_margin = p.margin
        section.left_margin   = p.left_margin
        section.right_margin  = p.margin
    normal = doc.styles["Normal"]
    normal.font.name                          = p.font
    normal.font.size                          = p.font_size
    normal.paragraph_format.line_spacing      = p.line_spacing
    normal.paragraph_format.space_before      = Pt(0)
    normal.paragraph_format.space_after       = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


# =============================================================================
# PARAGRAPH — tidak ada blank line, pakai space_after
# =============================================================================

def _add_para(doc: Document, p: StyleProfile, text: str = "",
              bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              first_indent=True, is_last_in_block=False):
    """
    Tambah paragraf teks.
    is_last_in_block=True → tambah para_space_after sebagai jarak ke elemen berikutnya.
    """
    para = doc.add_paragraph()
    para.alignment = align
    fmt = para.paragraph_format
    fmt.line_spacing      = p.line_spacing
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.space_before      = Pt(0)
    fmt.space_after       = p.para_space_after if is_last_in_block else Pt(0)
    if first_indent and p.first_line_indent > Pt(0):
        fmt.first_line_indent = p.first_line_indent
    else:
        fmt.first_line_indent = Pt(0)
    if text:
        run = para.add_run(text)
        run.font.name  = p.font
        run.font.size  = p.font_size
        run.bold       = bold
        run.italic     = italic
    return para


# =============================================================================
# HEADING — space_before/after sudah termasuk, tidak perlu blank paragraf
# =============================================================================

def _add_heading(doc: Document, p: StyleProfile, text: str, level: int = 1):
    para = doc.add_paragraph()
    fmt  = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing      = p.line_spacing
    fmt.first_line_indent = Pt(0)

    if level == 1:
        fmt.space_before = p.h1_space_before
        fmt.space_after  = p.h1_space_after
        para.alignment   = WD_ALIGN_PARAGRAPH.CENTER if p.h1_center else WD_ALIGN_PARAGRAPH.LEFT
        display          = text.upper() if p.h1_upper else text
        run = para.add_run(display)
        run.font.name = p.font; run.font.size = p.font_size
        run.bold = p.h1_bold; run.italic = False
        run.font.color.rgb = p.h1_color
    elif level == 2:
        fmt.space_before = p.h2_space_before
        fmt.space_after  = p.h2_space_after
        para.alignment   = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        run.font.name = p.font; run.font.size = p.font_size
        run.bold = p.h2_bold; run.italic = p.h2_italic
        run.font.color.rgb = p.h2_color
    else:  # level 3+
        fmt.space_before = p.h2_space_before
        fmt.space_after  = p.h2_space_after
        para.alignment   = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        run.font.name = p.font; run.font.size = p.font_size
        run.bold = True; run.italic = True
        run.font.color.rgb = p.h2_color
    return para


# =============================================================================
# TABLE HELPERS
# =============================================================================

def _remove_tbl_borders(tbl):
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    brd = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        brd.append(el)
    tbl_pr.append(brd)


def _set_cell_hline(cell, *, top=False, bottom=False, thick=False, color="000000"):
    tc    = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr"); tc.insert(0, tc_pr)
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)
    brd = OxmlElement("w:tcBorders")
    sz  = "12" if thick else "6"
    for side in ("top","bottom","left","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        if (side == "top" and top) or (side == "bottom" and bottom):
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "none")
        brd.append(el)
    tc_pr.append(brd)


def _set_cell_bg(cell, hex_color: str):
    tc    = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def _fmt_val(val) -> str:
    if isinstance(val, float):
        if abs(val) < 0.001 and val != 0:
            return "< .001"
        return f"{val:.3f}"
    return "" if val is None else str(val)


def _tbl_caption(doc: Document, p: StyleProfile, title: str, above: bool):
    """Render caption tabel. Tidak ada baris kosong — semua pakai space_before/after."""
    if not title:
        return
    parts = title.split(".", 1)
    lbl   = parts[0].strip()
    desc  = parts[1].strip() if len(parts) > 1 else ""

    p_lbl = doc.add_paragraph()
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_lbl.paragraph_format.line_spacing      = 1.0   # single di caption
    p_lbl.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_lbl.paragraph_format.space_before      = p.tbl_space_before if above else Pt(4)
    p_lbl.paragraph_format.space_after       = Pt(0)
    p_lbl.paragraph_format.first_line_indent = Pt(0)
    r = p_lbl.add_run(lbl)
    r.font.name = p.font; r.font.size = p.tbl_font_size
    r.bold = p.tbl_caption_bold_label

    if desc:
        p_dsc = doc.add_paragraph()
        p_dsc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_dsc.paragraph_format.line_spacing      = 1.0
        p_dsc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p_dsc.paragraph_format.space_before      = Pt(0)
        p_dsc.paragraph_format.space_after       = Pt(2)
        p_dsc.paragraph_format.first_line_indent = Pt(0)
        r2 = p_dsc.add_run(desc)
        r2.font.name = p.font; r2.font.size = p.tbl_font_size
        r2.italic = p.tbl_caption_italic_desc


def _style_table(doc: Document, dataframe: pd.DataFrame,
                 profile: StyleProfile, title: str = "") -> None:
    if dataframe is None or dataframe.empty:
        return

    if profile.tbl_caption_above and title:
        _tbl_caption(doc, profile, title, above=True)

    tbl = doc.add_table(rows=1, cols=len(dataframe.columns))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_tbl_borders(tbl)

    hdr_txt_rgb = RGBColor(
        int(profile.tbl_header_txt_hex[0:2], 16),
        int(profile.tbl_header_txt_hex[2:4], 16),
        int(profile.tbl_header_txt_hex[4:6], 16),
    )

    # Header
    for i, col_name in enumerate(dataframe.columns):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before      = Pt(3)
        para.paragraph_format.space_after       = Pt(3)
        para.paragraph_format.line_spacing      = 1.0
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.first_line_indent = Pt(0)
        run = para.add_run(str(col_name))
        run.font.name = profile.font; run.font.size = profile.tbl_font_size
        run.bold = True; run.font.color.rgb = hdr_txt_rgb

        if profile.tbl_has_color_header:
            _set_cell_bg(cell, profile.tbl_header_hex)
            _set_cell_hline(cell, top=True, bottom=True, thick=False,
                            color=profile.tbl_header_hex)
        else:
            _set_cell_hline(cell, top=True, bottom=True, thick=False, color="000000")
            tc_pr  = cell._tc.find(qn("w:tcPr"))
            tc_brd = tc_pr.find(qn("w:tcBorders"))
            if tc_brd is not None:
                top_el = tc_brd.find(qn("w:top"))
                if top_el is not None:
                    top_el.set(qn("w:sz"), "12")

    # Data rows
    rows_list = list(dataframe.iterrows())
    for row_idx, (_, row_data) in enumerate(rows_list):
        is_last  = (row_idx == len(rows_list) - 1)
        is_even  = (row_idx % 2 == 1)
        new_cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            cell = new_cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                              else WD_ALIGN_PARAGRAPH.CENTER)
            para.paragraph_format.space_before      = Pt(2)
            para.paragraph_format.space_after       = Pt(2)
            para.paragraph_format.line_spacing      = 1.0
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.first_line_indent = Pt(0)
            run = para.add_run(_fmt_val(val))
            run.font.name = profile.font; run.font.size = profile.tbl_font_size
            if profile.tbl_has_color_header and is_even:
                _set_cell_bg(cell, "F2F2F2")
            if is_last:
                _set_cell_hline(cell, bottom=True, thick=True, color="000000")
            else:
                _set_cell_hline(cell)

    # Spacer setelah tabel — space_after saja, BUKAN blank paragraf
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before      = Pt(0)
    sp.paragraph_format.space_after       = profile.tbl_space_after
    sp.paragraph_format.line_spacing      = 1.0
    sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    sp.paragraph_format.first_line_indent = Pt(0)

    if not profile.tbl_caption_above and title:
        _tbl_caption(doc, profile, title, above=False)


# =============================================================================
# GAMBAR / FIGURE
# =============================================================================

def _embed_image(doc: Document, profile: StyleProfile, png_bytes: bytes, label: str):
    try:
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts   = label.split(".", 1)
        fig_lbl = parts[0].strip()
        fig_dsc = parts[1].strip() if len(parts) > 1 else ""

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_cap.paragraph_format.space_before      = Pt(2)
        p_cap.paragraph_format.space_after       = profile.fig_space_after
        p_cap.paragraph_format.line_spacing      = 1.0
        p_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p_cap.paragraph_format.first_line_indent = Pt(0)
        r1 = p_cap.add_run(fig_lbl + ". ")
        r1.font.name = profile.font; r1.font.size = profile.tbl_font_size
        r1.bold   = profile.fig_caption_bold_label
        r1.italic = profile.fig_caption_italic_desc
        if fig_dsc:
            r2 = p_cap.add_run(fig_dsc)
            r2.font.name = profile.font; r2.font.size = profile.tbl_font_size
            r2.italic = profile.fig_caption_italic_desc
    except Exception as e:
        _add_para(doc, profile, f"[Gambar tidak dapat ditampilkan: {e}]",
                  first_indent=False)


# =============================================================================
# NARASI AI + FALLBACK OTOMATIS
# =============================================================================

def _clean_ai_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("**", "").replace("*", "")
    emoji_re = re.compile(
        "[\U00010000-\U0010ffff\U0001F600-\U0001F64F\U0001F300-\U0001F5FF"
        "\U0001F680-\U0001F6FF\U0001F1E0-\U0001F1FF\u2702-\u27B0\u24C2-\U0001F251]+",
        flags=re.UNICODE,
    )
    text = emoji_re.sub("", text)
    for sym in ["✓","✗","⚠️","❌","📐","🤖"]:
        text = text.replace(sym, "")
    return text.strip()


def _fallback_narasi(mod_key: str, data: dict) -> str:
    """
    Buat narasi interpretasi otomatis berstandar paper ilmiah dari data statistik,
    digunakan jika ai_texts[mod_key] kosong atau tidak tersedia.
    Seluruh narasi menggunakan diksi akademis formal sesuai konvensi jurnal ilmiah.
    """
    dep     = data.get("y", "variabel dependen")
    ind     = data.get("x", [])
    ind_str = ", ".join(ind) if ind else "variabel prediktor"

    # ── Interpretasi ukuran efek menggunakan standar Cohen (1988) ─────────────
    def _kat_d(v):
        v = abs(v)
        return "besar" if v >= 0.80 else ("sedang" if v >= 0.50 else "kecil")

    def _kat_eta2(v):
        return "besar" if v >= 0.14 else ("sedang" if v >= 0.06 else "kecil")

    def _kat_f2(v):
        return "besar" if v >= 0.35 else ("sedang" if v >= 0.15 else "kecil")

    def _kat_r(v):
        v = abs(v)
        return "kuat" if v >= 0.50 else ("sedang" if v >= 0.30 else
               "lemah" if v >= 0.10 else "sangat lemah")

    def _sig_str(pval, alpha=0.05):
        return "signifikan" if pval < alpha else "tidak signifikan"

    # ══════════════════════════════════════════════════════════════════════════
    # REGRESI LINIER & OLS+
    # ══════════════════════════════════════════════════════════════════════════
    if mod_key in ("regresi", "ols_plus"):
        r2      = data.get("r2")
        adj_r2  = data.get("adj_r2")
        pval    = data.get("f_pvalue")
        rmse    = data.get("rmse")
        dw      = data.get("durbin_watson")
        vif_max = data.get("vif_max")
        white_p = data.get("white_pvalue")
        shap_p  = data.get("shapiro_residual_p")

        r2_pct  = f"{r2 * 100:.1f}%" if r2 is not None else "—"
        adj_pct = f"{adj_r2 * 100:.1f}%" if adj_r2 is not None else "—"

        if r2 is not None:
            f2 = r2 / (1 - r2) if r2 < 1.0 else float("inf")
            kat_f2 = _kat_f2(f2)
        else:
            f2 = None
            kat_f2 = ""

        paragraf1 = (
            f"Analisis regresi linier berganda dilakukan untuk menguji pengaruh "
            f"{ind_str} terhadap {dep}. "
        )
        if pval is not None:
            sig_model = _sig_str(pval)
            paragraf1 += (
                f"Hasil uji F menunjukkan bahwa model secara simultan {sig_model} "
                f"(p = {pval:.3f}, \u03b1 = .05), "
            )
        if r2 is not None:
            paragraf1 += (
                f"dengan koefisien determinasi R\u00b2 = {r2:.3f} "
                f"(R\u00b2 Adjusted = {adj_pct}), "
                f"yang berarti model mampu menjelaskan {r2_pct} varians {dep}. "
            )
        if f2 is not None and not (f2 == float("inf")):
            paragraf1 += (
                f"Ukuran efek Cohen\u2019s f\u00b2 = {f2:.3f} termasuk kategori {kat_f2} "
                f"(Cohen, 1988), mengindikasikan relevansi praktis model yang "
                f"{'memadai' if f2 >= 0.15 else 'perlu ditingkatkan'}. "
            )

        # Paragraf koefisien
        lines_coef = []
        coef_df = data.get("coef_table")
        if coef_df is not None and not coef_df.empty:
            p_col   = next((c for c in coef_df.columns
                            if c.lower() in ("p", "p-value", "sig", "pval")), None)
            b_col   = next((c for c in coef_df.columns
                            if c.lower() in ("b", "coef", "koefisien", "\u03b2 (koefisien)",
                                             "\u03b2", "coefficient")), None)
            var_col = coef_df.columns[0]
            for _, row in coef_df.iterrows():
                var_name = str(row[var_col])
                if var_name.lower() in ("konstanta", "constant", "intercept"):
                    continue
                if p_col and b_col:
                    try:
                        pv = float(row[p_col])
                        bv = float(row[b_col])
                        ket_sig = (
                            "berpengaruh positif dan signifikan" if (bv > 0 and pv < 0.05) else
                            "berpengaruh negatif dan signifikan" if (bv < 0 and pv < 0.05) else
                            "tidak berpengaruh signifikan"
                        )
                        lines_coef.append(
                            f"Variabel {var_name} {ket_sig} terhadap {dep} "
                            f"(\u03b2 = {bv:.3f}, p = {pv:.3f})"
                        )
                    except Exception:
                        pass

        paragraf2 = ""
        if lines_coef:
            paragraf2 = (
                "Berdasarkan hasil pengujian koefisien regresi parsial, diperoleh temuan "
                "sebagai berikut: "
                + "; ".join(lines_coef) + ". "
                "Koefisien yang signifikan (p < .05) menunjukkan kontribusi unik variabel "
                "tersebut dalam memprediksi variabel dependen setelah mengontrol variabel "
                "lainnya (Field, 2018)."
            )

        # Paragraf asumsi klasik (khusus ols_plus)
        paragraf3 = ""
        if mod_key == "ols_plus":
            asumsi_lines = []
            if dw is not None:
                ket_dw = "terpenuhi" if 1.5 < float(dw) < 2.5 else "perlu diperiksa lebih lanjut"
                asumsi_lines.append(
                    f"autokorelasi diuji dengan Durbin-Watson = {float(dw):.3f} "
                    f"(asumsi {ket_dw})"
                )
            if vif_max is not None:
                ket_vif = ("tidak terdapat masalah multikolinearitas"
                           if float(vif_max) < 10
                           else "terdeteksi multikolinearitas (VIF > 10)")
                asumsi_lines.append(
                    f"VIF maksimum = {float(vif_max):.2f} ({ket_vif}, "
                    "acuan Marquardt, 1970)"
                )
            if white_p is not None:
                ket_het = ("homoskedastisitas terpenuhi"
                           if float(white_p) >= 0.05
                           else "terindikasi heteroskedastisitas")
                asumsi_lines.append(
                    f"uji White menunjukkan {ket_het} "
                    f"(p = {float(white_p):.3f}, White, 1980)"
                )
            if shap_p is not None:
                ket_norm = ("residual berdistribusi normal"
                            if float(shap_p) >= 0.05
                            else "residual tidak berdistribusi normal")
                asumsi_lines.append(
                    f"Shapiro-Wilk residual: {ket_norm} (p = {float(shap_p):.3f})"
                )
            if asumsi_lines:
                paragraf3 = (
                    "Evaluasi uji asumsi klasik menunjukkan: "
                    + "; ".join(asumsi_lines) + ". "
                    "Pemenuhan asumsi klasik merupakan prasyarat validitas inferensi "
                    "statistik dalam analisis regresi OLS (Hair et al., 2010)."
                )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # REGRESI LOGISTIK
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "logistik":
        auc     = data.get("auc")
        pseudo  = data.get("pseudo_r2")
        aic     = data.get("aic")
        bic     = data.get("bic")
        coef_df = _first_valid_df(data.get("coef_table"), data.get("odds_df"))

        paragraf1 = (
            f"Analisis regresi logistik biner dilakukan untuk memprediksi {dep} "
            f"berdasarkan {ind_str}. "
            "Model regresi logistik dipilih mengingat variabel dependen bersifat dikotomis, "
            "sehingga pendekatan ini lebih tepat dibandingkan regresi linier biasa "
            "(Hosmer & Lemeshow, 2013). "
        )
        if pseudo is not None:
            paragraf1 += (
                f"Pseudo R\u00b2 McFadden = {pseudo:.3f} mengindikasikan proporsi "
                "variansi yang dapat dijelaskan oleh model. "
            )
        if aic is not None and bic is not None:
            paragraf1 += f"Kriteria informasi model: AIC = {aic:.2f}, BIC = {bic:.2f}. "

        paragraf2 = ""
        if auc is not None:
            auc_f = float(auc)
            if auc_f >= 0.90:
                kat_auc, deskripsi = "sangat baik (excellent)", "sangat tinggi"
            elif auc_f >= 0.80:
                kat_auc, deskripsi = "baik (good)", "tinggi"
            elif auc_f >= 0.70:
                kat_auc, deskripsi = "cukup (fair)", "sedang"
            else:
                kat_auc, deskripsi = "rendah (poor)", "rendah"
            paragraf2 = (
                "Evaluasi performa diskriminatif model menggunakan area di bawah kurva ROC "
                f"(AUC) menghasilkan nilai {auc_f:.3f}, yang termasuk kategori {kat_auc}. "
                f"Nilai ini mengindikasikan kemampuan klasifikasi model yang {deskripsi} "
                f"dalam membedakan antara kedua kategori {dep} "
                "(Hosmer & Lemeshow, 2013). "
            )

        paragraf3 = ""
        or_lines = []
        if coef_df is not None and not coef_df.empty:
            p_col  = next((c for c in coef_df.columns
                           if c.lower() in ("p", "p-value", "sig")), None)
            or_col = next((c for c in coef_df.columns
                           if c.lower() in ("or (exp \u03b2)", "odds ratio",
                                            "exp(b)", "exp(coef)", "or")), None)
            var_col = coef_df.columns[0]
            if p_col and or_col:
                for _, row in coef_df.iterrows():
                    vn = str(row[var_col])
                    if vn.lower() in ("konstanta", "constant", "intercept"):
                        continue
                    try:
                        pv = float(row[p_col])
                        ov = float(row[or_col])
                        if pv < 0.05:
                            arah = "meningkatkan" if ov > 1 else "menurunkan"
                            or_lines.append(
                                f"{vn} secara signifikan {arah} kemungkinan {dep} "
                                f"sebesar {abs(ov - 1) * 100:.1f}% "
                                f"(OR = {ov:.3f}, p = {pv:.3f})"
                            )
                    except Exception:
                        pass
        if or_lines:
            paragraf3 = (
                "Berdasarkan nilai Odds Ratio (OR = exp[\u03b2]) variabel prediktor "
                "yang signifikan: "
                + "; ".join(or_lines) + ". "
                "OR > 1 menunjukkan peningkatan peluang kejadian, sedangkan OR < 1 "
                "menunjukkan penurunan peluang (Hosmer & Lemeshow, 2013)."
            )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # MEDIASI
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "mediasi":
        m        = data.get("m", "variabel mediator")
        x_var    = data.get("x", ind_str)
        boot_ci  = data.get("bootstrap_ci")
        indirect = data.get("indirect_effect")
        direct   = data.get("direct_effect")
        total    = data.get("total_effect")

        paragraf1 = (
            f"Analisis mediasi dilakukan untuk menguji peran {m} sebagai variabel "
            f"mediator dalam hubungan antara {x_var} dan {dep}, "
            "menggunakan prosedur bootstrapping sebagaimana direkomendasikan oleh "
            "Preacher dan Hayes (2008). "
        )

        paragraf2 = ""
        if boot_ci and indirect is not None:
            ci_lo, ci_hi = boot_ci
            tidak_nol = not (float(ci_lo) < 0 < float(ci_hi))
            sig_med = ("terdapat efek mediasi yang signifikan" if tidak_nol
                       else "tidak terdapat efek mediasi yang signifikan")
            paragraf2 = (
                f"Hasil bootstrapping (95% Bias-Corrected CI) menunjukkan {sig_med}, "
                f"dengan efek tidak langsung (indirect effect) = {float(indirect):.3f} "
                f"[{float(ci_lo):.3f}, {float(ci_hi):.3f}]. "
            )
            if tidak_nol and direct is not None and total is not None:
                d_f = float(direct)
                t_f = float(total)
                if abs(d_f) < 0.001:
                    jenis = "mediasi penuh (full mediation)"
                elif (d_f > 0) == (t_f > 0):
                    jenis = "mediasi sebagian (partial mediation)"
                else:
                    jenis = "mediasi suppressif"
                paragraf2 += (
                    f"Efek langsung (c\u2019) = {d_f:.3f}, efek total (c) = {t_f:.3f}, "
                    f"mengindikasikan {jenis} "
                    "(Baron & Kenny, 1986; Hayes, 2013). "
                )
            elif not tidak_nol:
                paragraf2 += (
                    "Interval kepercayaan yang mencakup nol mengindikasikan bahwa "
                    f"{m} tidak memediasi hubungan secara signifikan. "
                )

        paragraf3 = (
            f"Temuan ini memberikan pemahaman mekanistik mengenai bagaimana {x_var} "
            f"memengaruhi {dep} melalui {m} sebagai variabel perantara. "
            "Implikasi teoritis dari hasil mediasi ini perlu dikontekstualisasikan "
            "dalam kerangka teoretis penelitian yang lebih luas (Hayes, 2013)."
        )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # MODERASI
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "moderasi":
        z        = data.get("z", "variabel moderator")
        x_var    = data.get("x", ind_str)
        r2       = data.get("r2")
        adj_r2   = data.get("adj_r2")
        jn_point = data.get("johnson_neyman")

        paragraf1 = (
            f"Analisis moderasi dilakukan untuk menguji apakah {z} memoderasi "
            f"hubungan antara {x_var} dan {dep} "
            "(Aiken & West, 1991; Hayes, 2013). "
            "Pengujian dilakukan dengan memasukkan suku interaksi (X \u00d7 Z) ke dalam "
            "model regresi, di mana signifikansi suku interaksi menjadi indikator utama "
            "adanya efek moderasi. "
        )
        if r2 is not None:
            paragraf1 += (
                f"Model moderasi menjelaskan {r2 * 100:.1f}% variansi {dep} "
                f"(R\u00b2 = {r2:.3f}"
                + (f", R\u00b2 Adjusted = {adj_r2:.3f}" if adj_r2 is not None else "")
                + "). "
            )

        paragraf2 = ""
        coef_df = data.get("coef_table")
        interact_lines = []
        if coef_df is not None and not coef_df.empty:
            p_col   = next((c for c in coef_df.columns
                            if c.lower() in ("p", "p-value", "sig")), None)
            b_col   = next((c for c in coef_df.columns
                            if c.lower() in ("\u03b2", "b", "coef", "koefisien")), None)
            var_col = coef_df.columns[0]
            if p_col:
                for _, row in coef_df.iterrows():
                    vn = str(row[var_col])
                    if (":" in vn or "\u00d7" in vn or "*" in vn
                            or "interaction" in vn.lower()):
                        try:
                            pv = float(row[p_col])
                            bv = float(row[b_col]) if b_col else None
                            sig_int = _sig_str(pv)
                            ket = (
                                f"Suku interaksi {vn} {sig_int} (p = {pv:.3f}"
                                + (f", \u03b2 = {bv:.3f}" if bv is not None else "")
                                + "), "
                            )
                            if pv < 0.05:
                                ket += (
                                    f"menunjukkan bahwa {z} secara signifikan memoderasi "
                                    f"hubungan {x_var} \u2192 {dep}."
                                )
                            else:
                                ket += (
                                    f"menunjukkan bahwa {z} tidak terbukti memoderasi "
                                    f"hubungan {x_var} \u2192 {dep}."
                                )
                            interact_lines.append(ket)
                        except Exception:
                            pass
        if interact_lines:
            paragraf2 = " ".join(interact_lines)

        paragraf3 = ""
        if jn_point is not None:
            paragraf3 = (
                f"Analisis Johnson-Neyman mengidentifikasi titik transisi pada nilai "
                f"{z} = {float(jn_point):.3f}, di mana di bawah nilai tersebut hubungan "
                f"{x_var} \u2192 {dep} berubah signifikansinya (Johnson & Neyman, 1936). "
                "Temuan ini mengimplikasikan perlunya mempertimbangkan kondisi batas "
                "dalam perumusan rekomendasi praktis berbasis hasil penelitian ini."
            )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # ANOVA
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "anova":
        anova_df       = data.get("anova_table")
        eta2           = data.get("eta_squared")
        posthoc_df     = data.get("posthoc_table")
        posthoc_method = data.get("posthoc_method", "Tukey HSD")
        n_groups       = data.get("n_groups")
        num_col        = data.get("num_col", "variabel dependen")

        paragraf1 = (
            f"Uji analisis varians satu arah (One-Way ANOVA) dilakukan untuk menguji "
            f"perbedaan rata-rata {num_col} di antara kelompok yang dibandingkan. "
        )
        if n_groups:
            paragraf1 += f"Analisis melibatkan {n_groups} kelompok. "

        if anova_df is not None and not anova_df.empty:
            p_col = next((c for c in anova_df.columns
                          if c.lower() in ("p", "p-value", "pr(>f)", "p_value")), None)
            f_col = next((c for c in anova_df.columns
                          if c.lower() in ("f", "f-ratio", "f_stat")), None)
            if p_col:
                try:
                    pv = float(anova_df[p_col].iloc[0])
                    fv = float(anova_df[f_col].iloc[0]) if f_col else None
                    sig = _sig_str(pv)
                    paragraf1 += (
                        f"Hasil uji F"
                        + (f" = {fv:.3f} " if fv else " ")
                        + f"menunjukkan {sig} perbedaan rata-rata antar kelompok "
                        f"(p = {pv:.3f}). "
                    )
                    if pv < 0.05:
                        paragraf1 += (
                            "Dengan demikian, hipotesis nol yang menyatakan tidak terdapat "
                            "perbedaan rata-rata antar kelompok ditolak pada tingkat "
                            "signifikansi \u03b1 = .05. "
                        )
                    else:
                        paragraf1 += (
                            "Dengan demikian, tidak terdapat cukup bukti untuk menolak "
                            "hipotesis nol. "
                        )
                except Exception:
                    pass

        paragraf2 = ""
        if eta2 is not None:
            kat = _kat_eta2(float(eta2))
            paragraf2 = (
                f"Ukuran efek diukur menggunakan Eta Squared "
                f"(\u03b7\u00b2 = {float(eta2):.3f}), yang termasuk kategori {kat} "
                "berdasarkan konvensi Cohen (1988). "
                f"Nilai ini mengindikasikan bahwa {float(eta2) * 100:.1f}% variansi total "
                f"{num_col} dapat dijelaskan oleh perbedaan antar kelompok. "
                "Interpretasi ukuran efek penting untuk menilai signifikansi praktis "
                "di luar sekadar signifikansi statistik (Cohen, 1988; Field, 2018)."
            )

        paragraf3 = ""
        if posthoc_df is not None and not posthoc_df.empty:
            paragraf3 = (
                f"Untuk mengidentifikasi pasangan kelompok yang berbeda secara signifikan, "
                f"dilakukan uji post-hoc {posthoc_method}. "
                "Uji post-hoc berfungsi mengendalikan tingkat kesalahan tipe I pada "
                "perbandingan berganda (Tukey, 1949). "
                "Hasil selengkapnya dapat dilihat pada tabel post-hoc yang disajikan."
            )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # UJI BEDA
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "uji_beda":
        g1   = data.get("g1_name", "Kelompok 1")
        g2   = data.get("g2_name", "Kelompok 2")
        pval = data.get("p_value")
        stat = data.get("statistic")
        eff  = data.get("effect_size")
        uji  = data.get("uji_type", "t-test independen")
        m1   = data.get("g1_mean")
        m2   = data.get("g2_mean")

        paragraf1 = (
            f"Pengujian perbedaan dua kelompok dilakukan menggunakan {uji} "
            f"untuk membandingkan {g1} dan {g2}. "
        )
        if m1 is not None and m2 is not None:
            paragraf1 += (
                f"Rata-rata {g1} = {float(m1):.3f} "
                f"dan rata-rata {g2} = {float(m2):.3f}. "
            )

        paragraf2 = ""
        if pval is not None:
            sig = _sig_str(float(pval))
            ada = "terdapat" if float(pval) < 0.05 else "tidak terdapat"
            paragraf2 = (
                f"Hasil pengujian menunjukkan {ada} perbedaan yang {sig} secara statistik "
                f"antara {g1} dan {g2} "
                + (f"(statistik uji = {float(stat):.3f}, " if stat is not None else "(")
                + f"p = {float(pval):.3f}). "
                "Keputusan pengujian didasarkan pada tingkat signifikansi \u03b1 = .05 "
                "(Field, 2018; Cohen, 1988)."
            )

        paragraf3 = ""
        if eff is not None:
            kat = _kat_d(float(eff))
            paragraf3 = (
                f"Relevansi praktis perbedaan dinilai melalui ukuran efek "
                f"Cohen\u2019s d = {float(eff):.3f}, yang termasuk kategori {kat} "
                "menurut konvensi Cohen (1988). "
                "Ukuran efek yang signifikan secara statistik tidak selalu bermakna "
                "secara praktis, sehingga interpretasi Cohen\u2019s d menjadi penting "
                "untuk melengkapi keputusan inferensial."
            )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # SEM
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "sem":
        fit_df  = data.get("fit_indices")
        n_obs   = data.get("n_obs", "?")

        paragraf1 = (
            "Structural Equation Modeling (SEM) dilakukan secara terintegrasi dengan "
            "Confirmatory Factor Analysis (CFA) untuk menguji model pengukuran dan "
            "model struktural secara simultan. "
            f"Estimasi parameter dilakukan terhadap {n_obs} observasi. "
            "Pendekatan ini memungkinkan pengujian hubungan kausal antar variabel laten "
            "dengan mempertimbangkan kesalahan pengukuran "
            "(Hair et al., 2010; Hu & Bentler, 1999). "
        )

        paragraf2 = ""
        if fit_df is not None:
            try:
                is_df   = hasattr(fit_df, "empty")
                is_list = isinstance(fit_df, list)
                ok      = (is_df and not fit_df.empty) or (is_list and len(fit_df) > 0)
                if ok:
                    fit_data = fit_df if is_list else fit_df.to_dict("records")
                    idx_lines = []
                    thresholds = {
                        "cfi":   (0.90, "\u2265 .90", True),
                        "tli":   (0.90, "\u2265 .90", True),
                        "rmsea": (0.08, "\u2264 .08", False),
                        "srmr":  (0.08, "\u2264 .08", False),
                    }
                    for rec in fit_data:
                        if not isinstance(rec, dict):
                            continue
                        idx = str(rec.get("Indeks", rec.get("Index", ""))).lower()
                        val = rec.get("Nilai", rec.get("Value", rec.get("Hasil", "")))
                        for thr_key, (thr_val, thr_str, higher_better) in thresholds.items():
                            if thr_key in idx:
                                try:
                                    v  = float(str(val).replace("<", "").replace(">", ""))
                                    ok_fit = (v >= thr_val) if higher_better else (v <= thr_val)
                                    status = "memenuhi" if ok_fit else "tidak memenuhi"
                                    idx_lines.append(
                                        f"{thr_key.upper()} = {v:.3f} "
                                        f"({status} kriteria {thr_str})"
                                    )
                                except Exception:
                                    pass
                    if idx_lines:
                        paragraf2 = (
                            "Evaluasi kecocokan model (goodness-of-fit) menunjukkan: "
                            + "; ".join(idx_lines) + ". "
                            "Indeks kecocokan dievaluasi berdasarkan kriteria Hu dan Bentler "
                            "(1999): CFI dan TLI \u2265 .90, RMSEA \u2264 .08, SRMR \u2264 .08."
                        )
            except Exception:
                pass

        if not paragraf2:
            paragraf2 = (
                "Evaluasi kecocokan model dilakukan berdasarkan indeks CFI, TLI, RMSEA, "
                "dan SRMR (Hu & Bentler, 1999). "
                "Kriteria: CFI \u2265 .90, TLI \u2265 .90, RMSEA \u2264 .08, SRMR \u2264 .08. "
                "Interpretasi lengkap indeks kecocokan disajikan pada tabel di atas."
            )

        paragraf3 = (
            "Factor loadings pada model CFA mencerminkan validitas konvergen indikator "
            "terhadap konstruk laten. Loading \u2265 .50 dianggap cukup, sedangkan "
            "loading \u2265 .70 mengindikasikan validitas yang kuat (Hair et al., 2010). "
            "Estimasi jalur struktural dapat diinterpretasikan setelah memastikan "
            "kecocokan model yang memadai."
        )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # EFA — Analisis Faktor Eksploratori
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "efa":
        kmo      = data.get("kmo")
        kmo_lbl  = data.get("kmo_label", "")
        bart_p   = data.get("bartlett_p")
        n_fac    = data.get("n_factors")
        rotation = data.get("rotation", "Varimax")
        tot_v    = data.get("total_var")
        load_df  = data.get("loading_df")

        paragraf1 = (
            "Analisis Faktor Eksploratori (EFA) dilakukan untuk mengidentifikasi "
            "struktur laten yang mendasari seperangkat indikator pengukuran, "
            "dengan tujuan mereduksi dimensi data sekaligus mengungkap konstruk yang "
            "tidak dapat diobservasi secara langsung "
            "(Fabrigar et al., 1999; Hair et al., 2010). "
        )

        # Kelayakan data
        kel_lines = []
        if kmo is not None:
            kmo_f = float(kmo)
            kat_kmo = kmo_lbl if kmo_lbl else (
                "sangat memuaskan" if kmo_f >= 0.90 else
                "memuaskan"        if kmo_f >= 0.80 else
                "layak (middling)" if kmo_f >= 0.70 else
                "cukup (mediocre)" if kmo_f >= 0.60 else
                "tidak layak"
            )
            kel_lines.append(
                f"nilai Kaiser-Meyer-Olkin (KMO) = {kmo_f:.3f} ({kat_kmo}, Kaiser, 1974)"
            )
        if bart_p is not None:
            sig_bart = "signifikan" if float(bart_p) < 0.05 else "tidak signifikan"
            kel_lines.append(
                f"uji Bartlett of Sphericity {sig_bart} "
                f"(p = {float(bart_p):.4f}, Bartlett, 1950)"
            )

        paragraf2 = ""
        if kel_lines:
            layak = (
                (kmo is None or float(kmo) >= 0.60)
                and (bart_p is None or float(bart_p) < 0.05)
            )
            paragraf2 = (
                "Evaluasi kelayakan data untuk analisis faktor dilakukan melalui: "
                + " dan ".join(kel_lines) + ". "
                + (
                    "Data dinyatakan layak untuk dianalisis menggunakan EFA. "
                    if layak else
                    "Hasil uji mengindikasikan bahwa data perlu diperiksa kembali "
                    "sebelum dilanjutkan ke tahap analisis faktor. "
                )
            )

        paragraf3 = ""
        if n_fac is not None:
            paragraf3 = (
                f"Berdasarkan kriteria Kaiser (eigenvalue > 1) dan scree plot "
                f"(Cattell, 1966), diekstrak {n_fac} faktor menggunakan metode "
                f"rotasi {rotation}. "
            )
            if tot_v is not None:
                paragraf3 += (
                    f"Solusi faktor yang dihasilkan mampu menjelaskan "
                    f"{float(tot_v):.2f}% total variansi variabel-variabel yang dianalisis. "
                )

        # Loading summary
        paragraf4 = ""
        if load_df is not None and hasattr(load_df, "empty") and not load_df.empty:
            try:
                fac_cols = [
                    c for c in load_df.columns
                    if c.lower().startswith("factor")
                    or (c.lower().startswith("f") and c[1:].isdigit())
                ]
                if fac_cols:
                    n_high = 0
                    for col in fac_cols:
                        vals = pd.to_numeric(load_df[col], errors="coerce").abs()
                        n_high += int((vals >= 0.50).sum())
                    paragraf4 = (
                        f"Interpretasi faktor didasarkan pada nilai loading \u2265 .50 "
                        "sebagai kriteria item yang bermakna secara substantif "
                        "(Fabrigar et al., 1999). "
                        f"Sejumlah {n_high} item memiliki loading memadai pada faktor terkait. "
                        "Hasil EFA ini dapat menjadi dasar pengembangan instrumen lebih lanjut "
                        "melalui Confirmatory Factor Analysis (CFA) pada sampel independen "
                        "(Hair et al., 2010)."
                    )
            except Exception:
                pass

        if not paragraf4:
            paragraf4 = (
                "Interpretasi faktor dilakukan berdasarkan nilai factor loading setiap item, "
                "di mana loading \u2265 .50 dianggap bermakna secara substantif. "
                "Temuan EFA ini bersifat eksploratori dan perlu dikonfirmasi pada penelitian "
                "selanjutnya menggunakan CFA pada sampel yang berbeda "
                "(Fabrigar et al., 1999)."
            )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3, paragraf4] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # CFA — Confirmatory Factor Analysis
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "cfa":
        factor_map  = data.get("factor_map", {})
        n_konstruk  = len(factor_map)
        n_obs       = data.get("n_obs", "?")
        fit_records = data.get("fit_df_records", [])
        ave_records = data.get("ave_cr_df_records", [])

        paragraf1 = (
            "Confirmatory Factor Analysis (CFA) dilakukan untuk menguji validitas "
            f"konstruk model pengukuran yang terdiri dari {n_konstruk} konstruk laten "
            f"dengan {n_obs} observasi. "
            "CFA merupakan pendekatan berbasis teori yang menguji sejauh mana model "
            "pengukuran yang ditetapkan sesuai dengan data empiris "
            "(Hair et al., 2010; Hu & Bentler, 1999). "
            "Estimasi parameter menggunakan metode Maximum Likelihood (ML)."
        )

        paragraf2 = ""
        if fit_records:
            try:
                fit_df2 = pd.DataFrame(fit_records)
                if "Status" in fit_df2.columns:
                    n_good  = int(fit_df2["Status"].str.contains("\u2705").sum())
                    n_total = len(fit_df2)
                    pct     = n_good / n_total * 100
                    if n_good / n_total >= 0.75:
                        kat = "diterima (good fit)"
                    elif n_good / n_total >= 0.50:
                        kat = "diterima secara marginal (acceptable fit)"
                    else:
                        kat = ("tidak dapat diterima (poor fit) \u2014 "
                               "pertimbangkan respecifikasi model")
                    paragraf2 = (
                        f"Evaluasi kecocokan model menunjukkan {n_good} dari {n_total} "
                        f"indeks fit ({pct:.0f}%) memenuhi threshold yang ditetapkan, "
                        f"sehingga model secara keseluruhan dinyatakan {kat} "
                        "(Hu & Bentler, 1999). "
                        "Kriteria evaluasi meliputi CFI \u2265 .90, TLI \u2265 .90, "
                        "RMSEA \u2264 .08, dan SRMR \u2264 .08."
                    )
            except Exception:
                pass

        if not paragraf2:
            paragraf2 = (
                "Evaluasi kecocokan model dilakukan berdasarkan indeks CFI, TLI, "
                "RMSEA, dan SRMR (Hu & Bentler, 1999): "
                "CFI \u2265 .90, TLI \u2265 .90, RMSEA \u2264 .08, SRMR \u2264 .08."
            )

        paragraf3 = ""
        if ave_records:
            try:
                ave_df2 = pd.DataFrame(ave_records)
                ave_col = next(
                    (c for c in ave_df2.columns if "ave" in c.lower()), None
                )
                cr_col = next(
                    (c for c in ave_df2.columns
                     if "cr" in c.lower() or "composite" in c.lower()), None
                )
                if ave_col and cr_col:
                    n_ave = int(
                        (pd.to_numeric(ave_df2[ave_col], errors="coerce") >= 0.50).sum()
                    )
                    n_cr = int(
                        (pd.to_numeric(ave_df2[cr_col], errors="coerce") >= 0.70).sum()
                    )
                    n_tot = len(ave_df2)
                    paragraf3 = (
                        "Validitas konvergen dievaluasi melalui Average Variance Extracted "
                        f"(AVE) dan Composite Reliability (CR). "
                        f"AVE \u2265 .50 terpenuhi pada {n_ave}/{n_tot} konstruk, "
                        f"CR \u2265 .70 terpenuhi pada {n_cr}/{n_tot} konstruk "
                        "(Fornell & Larcker, 1981; Hair et al., 2010). "
                    )
                    if n_ave == n_tot and n_cr == n_tot:
                        paragraf3 += (
                            "Seluruh konstruk memenuhi kriteria validitas konvergen, "
                            "mengindikasikan bahwa indikator-indikator secara konsisten "
                            "mengukur konstruk laten yang dimaksud."
                        )
                    else:
                        paragraf3 += (
                            "Konstruk yang tidak memenuhi kriteria perlu ditinjau ulang, "
                            "baik melalui modifikasi model maupun penghapusan indikator "
                            "yang bermasalah."
                        )
            except Exception:
                pass

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # OUTLIER
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "outlier":
        n        = data.get("total_outliers", 0)
        n_total  = data.get("n_total", 0)
        pct      = data.get("pct_outliers", 0.0)
        variabel = data.get("variabel", "variabel yang dianalisis")
        method   = data.get("method", "IQR")

        if n > 0:
            return (
                f"Deteksi outlier pada variabel {variabel} menggunakan metode {method} "
                f"mengidentifikasi {n} observasi ({float(pct):.1f}% dari {n_total} total data) "
                "sebagai nilai pencilan (outlier). "
                "Keberadaan outlier berpotensi mendistorsi estimasi parameter regresi OLS "
                "maupun hasil uji statistik berbasis asumsi normalitas "
                "(Field, 2018; Ghozali, 2018). "
                "Penanganan outlier harus dilakukan berdasarkan pertimbangan substantif: "
                "jika outlier merupakan kesalahan pengukuran atau entri data, penghapusan "
                "dapat dibenarkan; namun jika mencerminkan variasi alami populasi, outlier "
                "sebaiknya dipertahankan menggunakan metode estimasi yang lebih robust. "
                "Analisis sensitivitas (dengan dan tanpa outlier) direkomendasikan untuk "
                "menilai pengaruhnya terhadap kesimpulan penelitian."
            )
        else:
            return (
                f"Hasil deteksi outlier pada variabel {variabel} menggunakan metode {method} "
                f"tidak menemukan nilai pencilan dari {n_total} observasi yang diperiksa. "
                "Seluruh data berada dalam rentang yang diharapkan berdasarkan distribusi "
                "empiris variabel, sehingga dinyatakan layak untuk dianalisis lebih lanjut "
                "tanpa perlakuan khusus."
            )

    # ══════════════════════════════════════════════════════════════════════════
    # ANALISIS KELOMPOK
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "kelompok":
        cat   = data.get("cat", "kelompok")
        num   = data.get("num", "variabel")
        pval  = data.get("p_value")
        fstat = data.get("f_stat")
        best  = data.get("best_group")
        worst = data.get("worst_group")

        paragraf1 = (
            f"Analisis perbandingan variabel \u2018{num}\u2019 berdasarkan kategori "
            f"\u2018{cat}\u2019 dilakukan untuk mengidentifikasi perbedaan karakteristik "
            "antar kelompok secara statistik (Field, 2018). "
        )

        paragraf2 = ""
        if pval is not None:
            ada = "terdapat" if float(pval) < 0.05 else "tidak terdapat"
            sig = _sig_str(float(pval))
            paragraf2 = (
                f"Hasil pengujian menunjukkan {ada} perbedaan yang {sig} "
                f"pada variabel {num} antar kelompok {cat} "
                + (f"(F = {float(fstat):.3f}, " if fstat else "(")
                + f"p = {float(pval):.3f}). "
            )
            if best and float(pval) < 0.05:
                paragraf2 += (
                    f"Kelompok {best} memiliki rata-rata tertinggi"
                    + (f", sedangkan kelompok {worst} memiliki rata-rata terendah. "
                       if worst else ". ")
                )

        paragraf3 = (
            "Temuan ini memberikan implikasi bagi pengambilan keputusan berbasis kelompok. "
            "Perbedaan yang signifikan mengindikasikan perlunya strategi atau intervensi "
            "yang disesuaikan dengan karakteristik masing-masing kelompok."
        )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # RELIABILITAS ICC
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "reliabilitas_icc":
        rec_model   = data.get("rec_model", "ICC(2,1)")
        n_subj      = data.get("n_subj", "?")
        n_rater     = data.get("n_rater", "?")
        use_type    = data.get("use_type", "rater agreement")
        icc_records = data.get("icc_df", [])

        paragraf1 = (
            "Uji reliabilitas menggunakan Intraclass Correlation Coefficient (ICC) "
            "dilakukan untuk mengevaluasi tingkat konsistensi dan keandalan pengukuran "
            f"dalam konteks {use_type}, "
            f"melibatkan {n_subj} subjek dan {n_rater} rater atau sesi pengukuran. "
            "ICC merupakan indeks yang digunakan secara luas untuk menilai reliabilitas "
            "antar-penilai maupun tes-retes pada data pengukuran berulang "
            "(Koo & Mae, 2016; Shrout & Fleiss, 1979). "
        )

        paragraf2 = ""
        if icc_records:
            try:
                for row in icc_records:
                    if isinstance(row, dict) and row.get("Model") == rec_model:
                        icc_val = row.get("ICC")
                        p_val   = row.get("p_value")
                        ci_lo   = row.get("CI_Lower")
                        ci_hi   = row.get("CI_Upper")
                        if icc_val is not None:
                            v = float(icc_val)
                            if v >= 0.90:
                                kat = "sangat baik (excellent)"
                                imp = ("pengukuran dapat diandalkan untuk pengambilan "
                                       "keputusan klinis maupun penelitian")
                            elif v >= 0.75:
                                kat = "baik (good)"
                                imp = ("pengukuran memiliki konsistensi yang cukup "
                                       "untuk keperluan penelitian")
                            elif v >= 0.50:
                                kat = "sedang (moderate)"
                                imp = ("pengukuran perlu ditingkatkan sebelum "
                                       "digunakan secara luas")
                            else:
                                kat = "buruk (poor)"
                                imp = "prosedur pengukuran perlu direvisi secara substansial"
                            paragraf2 = (
                                f"Model ICC yang direkomendasikan ({rec_model}) "
                                f"menghasilkan ICC = {v:.4f} ({kat})"
                                + (f", 95% CI [{float(ci_lo):.4f}, {float(ci_hi):.4f}]"
                                   if ci_lo is not None and ci_hi is not None else "")
                                + (f", p = {float(p_val):.4f}" if p_val is not None else "")
                                + f". Nilai ini mengindikasikan bahwa {imp} "
                                  "(Koo & Mae, 2016). "
                            )
                        break
            except Exception:
                pass

        if not paragraf2:
            paragraf2 = (
                "Interpretasi ICC mengikuti kriteria Koo dan Mae (2016): "
                "ICC < .50 = buruk, .50\u2013.75 = sedang, .75\u2013.90 = baik, "
                "\u2265 .90 = sangat baik. "
                "Nilai ICC yang tinggi mengindikasikan konsistensi pengukuran yang "
                "dapat diandalkan antar rater maupun antar waktu pengukuran."
            )

        paragraf3 = (
            "Pemilihan model ICC yang tepat bergantung pada desain studi: "
            "ICC(1) untuk rater acak tanpa efek campuran; "
            "ICC(2) untuk rater yang dipilih dan digeneralisasikan ke populasi; "
            "ICC(3) untuk rater yang dipilih tanpa generalisasi "
            "(Shrout & Fleiss, 1979). "
            "Perbedaan antara \u2018agreement\u2019 dan \u2018consistency\u2019 perlu "
            "diperhatikan: agreement memperhitungkan perbedaan sistematis antar rater, "
            "sedangkan consistency hanya menilai konsistensi relatif."
        )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # UJI ASUMSI PRA-ANALISIS
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "uji_asumsi":
        rec = data.get("rekomendasi", {})
        if not isinstance(rec, dict):
            return (
                "Uji asumsi pra-analisis telah dilakukan untuk menentukan pendekatan "
                "statistik yang paling sesuai dengan karakteristik data. "
                "Hasil uji asumsi menjadi dasar pemilihan metode analisis, apakah "
                "parametrik atau non-parametrik (Field, 2018)."
            )

        level  = rec.get("level", "")
        skor   = rec.get("skor_lulus", "?")
        total  = rec.get("total_uji", "?")
        pct    = rec.get("pct_lulus", "?")
        detail = rec.get("detail", {})

        paragraf1 = (
            "Serangkaian uji asumsi statistik dilakukan sebagai prosedur pra-analisis "
            "untuk menentukan kesesuaian data terhadap persyaratan metode statistik parametrik. "
            f"Hasil evaluasi menunjukkan {skor} dari {total} uji asumsi terpenuhi "
            f"({pct}%)"
            + (f", sehingga data direkomendasikan untuk analisis {level}. "
               if level else ". ")
        )

        paragraf2 = ""
        if isinstance(detail, dict) and detail:
            asumsi_lines = []
            labels = {
                "normalitas_univariat":   ("Normalitas univariat (Shapiro-Wilk)",
                                           "n_normal", "n_variabel"),
                "normalitas_multivariat": ("Normalitas multivariat (Mardia\u2019s test)",
                                           "p_skew", "p_kurt"),
                "homogenitas":            ("Homogenitas varians (Levene & Bartlett)",
                                           "n_lulus", "n_uji"),
                "linieritas":             ("Linieritas (Ramsey RESET test)",
                                           "n_lulus", "n_uji"),
            }
            for key, info in detail.items():
                if not isinstance(info, dict):
                    continue
                lbl_tuple = labels.get(key)
                lbl   = lbl_tuple[0] if lbl_tuple else key
                lulus = info.get("lulus", False)
                status = "terpenuhi" if lulus else "tidak terpenuhi"
                if key == "normalitas_univariat":
                    n_n = info.get("n_normal", "?")
                    n_v = info.get("n_variabel", "?")
                    ket = f"{n_n}/{n_v} variabel berdistribusi normal"
                elif key == "normalitas_multivariat":
                    ps = info.get("p_skew", "?")
                    pk = info.get("p_kurt", "?")
                    ket = f"skewness p = {ps}, kurtosis p = {pk}"
                elif key == "homogenitas":
                    nl = info.get("n_lulus", "?")
                    nu = info.get("n_uji", "?")
                    ket = f"{nl}/{nu} kelompok homogen"
                elif key == "linieritas":
                    nl = info.get("n_lulus", "?")
                    nu = info.get("n_uji", "?")
                    ket = f"{nl}/{nu} pasangan memenuhi asumsi linieritas"
                else:
                    ket = ""
                asumsi_lines.append(
                    f"{lbl}: {status} ({ket})" if ket else f"{lbl}: {status}"
                )
            if asumsi_lines:
                paragraf2 = (
                    "Rincian hasil uji asumsi per komponen: "
                    + "; ".join(asumsi_lines) + ". "
                )

        paragraf3 = (
            "Pemenuhan asumsi statistik merupakan prasyarat validitas inferensi pada "
            "metode parametrik. Apabila asumsi normalitas tidak terpenuhi, penggunaan "
            "uji non-parametrik (Mann-Whitney, Kruskal-Wallis) atau transformasi data "
            "direkomendasikan. Pelanggaran homogenitas varians dapat diatasi dengan "
            "uji Welch atau estimasi robust (Field, 2018)."
        )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # REGRESI ROBUST
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "ols_robust":
        dep_var   = data.get("dep_var", "variabel dependen")
        ind_vars  = data.get("ind_vars", [])
        ind_s     = ", ".join(ind_vars) if ind_vars else "variabel prediktor"
        estimator = data.get("estimator", "Huber-M")
        n_low     = data.get("n_low_weight", 0)
        n_obs     = data.get("n_obs", "?")
        n_changed = data.get("n_changed", 0)
        ols_rmse  = data.get("ols_rmse")
        rlm_rmse  = data.get("rlm_rmse")

        paragraf1 = (
            f"Regresi Robust menggunakan estimator M ({estimator}) dilakukan untuk "
            f"memodelkan {dep_var} berdasarkan {ind_s} (N = {n_obs}). "
            "Pendekatan ini dipilih karena estimator OLS sensitif terhadap outlier dan "
            "leverage points yang dapat mendistorsi estimasi koefisien secara substansial. "
            "Regresi Robust mereduksi pengaruh observasi bermasalah dengan memberikan "
            "bobot lebih rendah secara iteratif (Huber, 1973; Greene, 2012)."
        )

        p2_lines = []
        if n_low:
            p2_lines.append(
                f"sebanyak {n_low} observasi memperoleh bobot rendah (< 0.5), "
                "mengindikasikan keberadaan outlier atau leverage points"
            )
        if n_changed:
            p2_lines.append(
                f"{n_changed} koefisien mengalami perubahan > 10% dibandingkan OLS, "
                "menunjukkan sensitivitas OLS terhadap data bermasalah"
            )
        if ols_rmse is not None and rlm_rmse is not None:
            rlm_f = float(rlm_rmse)
            ols_f = float(ols_rmse)
            lebih = "lebih rendah" if rlm_f < ols_f else "sebanding"
            p2_lines.append(
                f"RMSE Robust ({rlm_f:.4f}) {lebih} dibandingkan OLS ({ols_f:.4f})"
            )

        paragraf2 = ""
        if p2_lines:
            paragraf2 = (
                "Hasil diagnostik regresi robust menunjukkan: "
                + "; ".join(p2_lines) + ". "
                "Perbandingan koefisien OLS dan Robust disarankan untuk menilai "
                "robustness temuan penelitian (Huber, 1973)."
            )

        return "\n\n".join(p for p in [paragraf1, paragraf2] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # WLS
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "ols_wls":
        dep_var       = data.get("dep_var", "variabel dependen")
        ind_vars      = data.get("ind_vars", [])
        ind_s         = ", ".join(ind_vars) if ind_vars else "variabel prediktor"
        weight_method = data.get("weight_method", "1/|\u03b5|")
        n_obs         = data.get("n_obs", "?")
        ols_g         = data.get("ols_glejser_p")
        wls_g         = data.get("wls_glejser_p")
        ols_r         = data.get("ols_rmse")
        wls_r         = data.get("wls_rmse")
        n_changed     = data.get("n_changed", 0)

        paragraf1 = (
            f"Weighted Least Squares (WLS) diterapkan untuk mengatasi masalah "
            f"heteroskedastisitas pada model {dep_var} ~ {ind_s} (N = {n_obs}), "
            f"menggunakan skema pembobotan {weight_method}. "
            "Heteroskedastisitas melanggar asumsi OLS dan menghasilkan estimasi yang "
            "tidak efisien, sehingga inferensi berbasis standard error OLS menjadi "
            "tidak valid (Greene, 2012; White, 1980). "
        )

        p2_lines = []
        if ols_g is not None and wls_g is not None:
            ols_gf = float(ols_g)
            wls_gf = float(wls_g)
            if ols_gf < 0.05 and wls_gf >= 0.05:
                status = "WLS berhasil mengatasi heteroskedastisitas"
            elif ols_gf < 0.05 and wls_gf < 0.05:
                status = ("heteroskedastisitas masih terdeteksi setelah WLS "
                          "\u2014 pertimbangkan transformasi tambahan")
            else:
                status = ("OLS tidak menunjukkan heteroskedastisitas signifikan, "
                          "WLS sebagai robustness check")
            p2_lines.append(
                f"uji Glejser: OLS p = {ols_gf:.4f} \u2192 WLS p = {wls_gf:.4f} "
                f"({status})"
            )
        if ols_r is not None and wls_r is not None:
            wls_rf = float(wls_r)
            ols_rf = float(ols_r)
            p2_lines.append(
                f"RMSE: OLS = {ols_rf:.4f}, WLS = {wls_rf:.4f} "
                f"({'WLS lebih efisien' if wls_rf < ols_rf else 'OLS sebanding'})"
            )
        if n_changed:
            p2_lines.append(f"{n_changed} koefisien berubah > 10% dari OLS ke WLS")

        paragraf2 = ""
        if p2_lines:
            paragraf2 = (
                "Perbandingan diagnostik OLS versus WLS: "
                + "; ".join(p2_lines) + ". "
                "Koefisien WLS yang berbeda signifikan dari OLS mengindikasikan bahwa "
                "heteroskedastisitas berpengaruh pada estimasi (Greene, 2012)."
            )

        return "\n\n".join(p for p in [paragraf1, paragraf2] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # PERBANDINGAN MODEL REGRESI
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "ols_robust_comparison":
        dep_var    = data.get("dep_var", "variabel dependen")
        ind_vars   = data.get("ind_vars", [])
        ind_s      = ", ".join(ind_vars) if ind_vars else "variabel prediktor"
        best_model = data.get("best_model", "OLS (Baseline)")
        n_obs      = data.get("n_obs", "?")

        paragraf1 = (
            "Perbandingan empat pendekatan regresi \u2014 OLS (baseline), RLM Huber-M, "
            f"RLM Bisquare, dan WLS \u2014 dilakukan untuk memilih estimator yang paling "
            f"tepat dalam memodelkan {dep_var} ~ {ind_s} (N = {n_obs}). "
            "Pendekatan komparatif ini memastikan bahwa kesimpulan penelitian bersifat "
            "robust terhadap berbagai kondisi data (Greene, 2012; Huber, 1973)."
        )

        paragraf2 = (
            f"Berdasarkan kriteria RMSE sebagai ukuran akurasi prediksi, model terbaik "
            f"adalah {best_model}. "
            "RMSE yang lebih rendah mengindikasikan deviasi prediksi yang lebih kecil, "
            "sehingga model tersebut direkomendasikan untuk analisis lanjutan. "
            "Perbandingan koefisien antar model disarankan untuk menilai stabilitas "
            "estimasi dan mengidentifikasi pengaruh observasi bermasalah."
        )

        return "\n\n".join(p for p in [paragraf1, paragraf2] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # COMPUTE VARIABEL
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "compute":
        compute_log = data.get("compute_log", [])
        n_ops       = len(compute_log)

        if n_ops == 0:
            return (
                "Tidak terdapat operasi komputasi variabel yang tercatat dalam sesi ini. "
                "Transformasi atau pembentukan variabel baru dapat dilakukan melalui modul "
                "Compute Variabel sebelum analisis statistik utama dijalankan."
            )

        methods  = list({e.get("method", "") for e in compute_log if e.get("method")})
        new_cols = [e.get("new_col", "") for e in compute_log[:8] if e.get("new_col")]

        paragraf1 = (
            f"Sebanyak {n_ops} variabel baru dibentuk melalui operasi komputasi "
            f"yang mencakup: {', '.join(methods) if methods else 'berbagai metode transformasi'}. "
            "Pembentukan variabel baru merupakan bagian dari tahap persiapan data yang "
            "bertujuan menghasilkan indikator pengukuran yang sesuai dengan kerangka "
            "konseptual penelitian (Field, 2018). "
        )

        paragraf2 = ""
        if new_cols:
            paragraf2 = (
                f"Variabel baru yang dibentuk antara lain: {', '.join(new_cols)}"
                + (" (dan lainnya)" if n_ops > 8 else "")
                + ". Setiap variabel komposit atau transformasi harus didokumentasikan "
                "secara eksplisit dalam laporan metodologi penelitian, termasuk dasar "
                "teoritis dan prosedur komputasi yang digunakan."
            )

        return "\n\n".join(p for p in [paragraf1, paragraf2] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # ANALISIS KLASTER
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "klaster":
        method     = data.get("method", "K-Means")
        k          = data.get("k", "?")
        cols_kl    = data.get("cols", [])
        silhouette = data.get("silhouette")
        linkage    = data.get("linkage", "ward")
        n_var      = len(cols_kl)

        paragraf1 = (
            f"Analisis klaster dilakukan menggunakan metode {method} untuk "
            f"mengelompokkan observasi ke dalam {k} klaster berdasarkan {n_var} variabel"
            + (f": {', '.join(cols_kl[:6])}"
               + (" (dan lainnya)" if n_var > 6 else "")
               if cols_kl else "")
            + ". Data distandarisasi menggunakan Z-score sebelum proses clustering "
            "untuk menghilangkan pengaruh perbedaan skala pengukuran "
            "(MacQueen, 1967; Ward, 1963). "
        )
        if method == "Hierarchical":
            paragraf1 += (
                f"Metode hierarchical clustering dengan linkage {linkage} digunakan, "
                "di mana jumlah klaster optimal ditentukan berdasarkan inspeksi "
                "dendrogram."
            )

        paragraf2 = ""
        if silhouette is not None:
            sv = float(silhouette)
            if sv >= 0.70:
                kat  = "sangat baik"
                interp = ("klaster-klaster terbentuk dengan kohesi internal yang tinggi "
                          "dan separasi yang jelas")
            elif sv >= 0.50:
                kat  = "baik"
                interp = "klaster terbentuk dengan kohesi internal yang memadai"
            elif sv >= 0.30:
                kat  = "cukup (acceptable)"
                interp = "klaster terbentuk namun dengan overlap yang perlu diperhatikan"
            else:
                kat  = "rendah"
                interp = ("kualitas klaster perlu ditingkatkan \u2014 pertimbangkan "
                          "jumlah klaster yang berbeda")
            paragraf2 = (
                f"Kualitas pengelompokan dievaluasi menggunakan Silhouette Score = {sv:.4f}, "
                f"yang termasuk kategori {kat} (Rousseeuw, 1987). "
                f"Nilai ini mengindikasikan bahwa {interp}. "
                "Interpretasi profil klaster dilakukan dengan membandingkan rata-rata "
                "masing-masing variabel antar klaster untuk merumuskan karakteristik "
                "dominan setiap segmen."
            )

        paragraf3 = (
            "Hasil analisis klaster bersifat deskriptif-eksploratif dan tidak menghasilkan "
            "inferensi statistik seperti pada uji hipotesis. Validitas temuan klaster dapat "
            "ditingkatkan melalui prosedur cross-validation atau replikasi pada sampel "
            "independen (Rousseeuw, 1987)."
        )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # EDA
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "eda":
        n_rows   = data.get("n_rows", "?")
        n_cols   = data.get("n_cols", "?")
        n_num    = data.get("n_numeric", "?")
        n_cat    = data.get("n_cat", "?")
        n_miss   = data.get("n_missing", 0)
        pct_miss = data.get("pct_missing", 0)
        n_dup    = data.get("n_dup", 0)

        paragraf1 = (
            "Eksplorasi Data Awal (Exploratory Data Analysis/EDA) dilakukan terhadap "
            f"dataset yang terdiri dari {n_rows} observasi dan {n_cols} variabel "
            f"({n_num} numerik, {n_cat} kategorik). "
            "EDA merupakan tahap fundamental yang bertujuan memahami karakteristik "
            "distribusi data, mengidentifikasi anomali, dan menentukan strategi analisis "
            "yang tepat sebelum pengujian hipotesis dilakukan (Tukey, 1977; Field, 2018)."
        )

        p2_lines = []
        try:
            pct_f = float(pct_miss)
            mi    = int(n_miss)
            di    = int(n_dup)
            if mi > 0:
                if pct_f > 20:
                    p2_lines.append(
                        f"missing values tinggi ({mi} sel, {pct_f:.1f}%) \u2014 "
                        "direkomendasikan multiple imputation atau penghapusan variabel "
                        "dengan > 30% missing"
                    )
                elif pct_f > 5:
                    p2_lines.append(
                        f"missing values moderat ({mi} sel, {pct_f:.1f}%) \u2014 "
                        "direkomendasikan imputasi (mean, median, atau MICE)"
                    )
                else:
                    p2_lines.append(
                        f"missing values minimal ({mi} sel, {pct_f:.1f}%) \u2014 "
                        "dapat ditangani dengan listwise deletion atau imputasi sederhana"
                    )
            if di > 0:
                p2_lines.append(
                    f"{di} baris duplikat terdeteksi \u2014 perlu dievaluasi apakah "
                    "merupakan entri ganda atau data sah"
                )
        except Exception:
            pass

        paragraf2 = ""
        if p2_lines:
            paragraf2 = (
                "Evaluasi kualitas data mengidentifikasi: "
                + "; ".join(p2_lines) + ". "
                "Penanganan yang tepat diperlukan sebelum analisis statistik inferensial "
                "dilanjutkan."
            )
        else:
            try:
                if int(n_miss) == 0 and int(n_dup) == 0:
                    paragraf2 = (
                        "Dataset berada dalam kondisi baik: tidak terdapat missing values "
                        "maupun baris duplikat. Data siap dianalisis lebih lanjut tanpa "
                        "memerlukan prosedur pembersihan tambahan."
                    )
            except Exception:
                pass

        return "\n\n".join(p for p in [paragraf1, paragraf2] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # SCRAPING
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "scraping":
        source = data.get("source", "sumber tidak diketahui")
        n_rows = data.get("n_rows", "?")
        n_cols = data.get("n_cols", "?")
        n_num  = data.get("n_numeric", "?")
        n_miss = data.get("n_missing", 0)
        n_dup  = data.get("n_dup", 0)

        paragraf1 = (
            "Data penelitian diperoleh secara otomatis melalui teknik web scraping dari "
            f"sumber: {source}. Dataset yang berhasil dikumpulkan terdiri dari {n_rows} "
            f"observasi dan {n_cols} variabel ({n_num} numerik). "
            "Pengambilan data dilakukan dengan memperhatikan kebijakan robots.txt dan "
            "rate-limiting untuk mematuhi etika pengambilan data publik digital."
        )

        kl = []
        try:
            if int(n_miss) > 0:
                kl.append(f"{n_miss} missing values yang ditangani melalui cleaning")
            if int(n_dup) > 0:
                kl.append(f"{n_dup} baris duplikat yang dieliminasi pada preprocessing")
        except Exception:
            pass

        paragraf2 = (
            "Proses data cleaning meliputi: " + "; ".join(kl) + ". "
            "Standardisasi format, konversi tipe variabel, dan validasi rentang nilai "
            "dilakukan sebelum analisis untuk memastikan integritas dataset."
            if kl else
            "Dataset hasil scraping tidak memiliki missing values maupun duplikat, "
            "sehingga siap dianalisis setelah validasi format dan tipe data."
        )

        paragraf3 = (
            "Validitas data sekunder yang dikumpulkan melalui scraping perlu dikritisi "
            "mengingat potensi perubahan struktur halaman web, bias seleksi sumber, dan "
            "keterbatasan aksesibilitas. Reproducibility penelitian dapat ditingkatkan "
            "dengan menyimpan log scraping yang mencatat waktu pengambilan, URL, dan "
            "versi dataset yang digunakan."
        )

        return "\n\n".join(p for p in [paragraf1, paragraf2, paragraf3] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # POWER ANALYSIS
    # ══════════════════════════════════════════════════════════════════════════
    elif mod_key == "power_analysis":
        test_type   = data.get("test_type", "uji statistik")
        effect_size = data.get("effect_size")
        alpha       = data.get("alpha", 0.05)
        power       = data.get("power", 0.80)
        n_total     = data.get("n_total")
        n_per_group = data.get("n_per_group")

        paragraf1 = (
            f"Analisis kekuatan statistik (power analysis) dilakukan untuk {test_type} "
            "guna menentukan ukuran sampel minimum yang diperlukan agar penelitian "
            "memiliki kekuatan statistik yang memadai. "
            f"Tingkat signifikansi yang ditetapkan adalah \u03b1 = {alpha}, "
            f"dengan target statistical power = {power} (Cohen, 1988). "
        )

        paragraf2 = ""
        if effect_size is not None:
            paragraf2 = (
                f"Berdasarkan ukuran efek yang diantisipasi (effect size = "
                f"{float(effect_size):.3f}) dan tingkat power {power}, "
            )
            if n_total is not None:
                paragraf2 += f"diperlukan total sampel minimum N = {n_total}"
                if n_per_group is not None:
                    paragraf2 += f" ({n_per_group} per kelompok)"
                paragraf2 += ". "
            paragraf2 += (
                "Ukuran sampel yang memadai memastikan bahwa penelitian memiliki "
                "kemampuan yang cukup untuk mendeteksi efek yang sesungguhnya ada di "
                "populasi (Cohen, 1988; Field, 2018)."
            )

        return "\n\n".join(p for p in [paragraf1, paragraf2] if p.strip())

    # ══════════════════════════════════════════════════════════════════════════
    # Fallback generik
    # ══════════════════════════════════════════════════════════════════════════
    return (
        "Analisis statistik telah dilakukan sesuai dengan prosedur yang ditetapkan. "
        "Interpretasi hasil disajikan pada tabel dan visualisasi yang menyertai bagian ini. "
        "Pembaca disarankan merujuk pada nilai statistik uji, tingkat signifikansi, "
        "dan ukuran efek yang dilaporkan untuk menarik kesimpulan yang tepat "
        "(Field, 2018; Cohen, 1988)."
    )

def _add_ai_narasi(doc: Document, profile: StyleProfile,
                   ai_text: str, title: str,
                   mod_key: str = "", data: dict = None):
    """
    Tambah narasi interpretasi.
    Jika ai_text kosong, gunakan _fallback_narasi() dari data statistik.
    """
    clean = ""
    if ai_text and not str(ai_text).startswith(("❌","⚠️")):
        clean = _clean_ai_text(ai_text)

    # Fallback ke narasi otomatis dari data
    if not clean and mod_key and data:
        try:
            clean = _fallback_narasi(mod_key, data)
        except Exception:
            clean = ""

    if not clean:
        return

    if profile.ai_show_label:
        _add_heading(doc, profile, title, level=profile.ai_label_level)

    paras = [l.strip() for l in clean.split("\n") if l.strip()]
    for i, para_text in enumerate(paras):
        _add_para(doc, profile, para_text, first_indent=True,
                  is_last_in_block=(i == len(paras) - 1))


def _add_model_equation_box(doc: Document, profile: StyleProfile, equation_text: str):
    if not equation_text or str(equation_text).startswith(("❌","⚠️")):
        return
    clean = _clean_ai_text(equation_text)
    if not clean:
        return
    _add_heading(doc, profile, "Model Persamaan", level=3)
    indent = profile.first_line_indent if profile.first_line_indent > Pt(0) else Cm(1.0)
    for line in [l.strip() for l in clean.split("\n") if l.strip()]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = p.paragraph_format
        fmt.left_indent       = indent
        fmt.first_line_indent = Pt(0)
        fmt.space_before      = Pt(0)
        fmt.space_after       = Pt(0)
        fmt.line_spacing      = profile.line_spacing
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        run = p.add_run(line)
        run.font.name = profile.font; run.font.size = profile.font_size
        run.italic = True


# =============================================================================
# MODULE RENDERERS
# =============================================================================

def _get_eq(analysis_type, data, ai_texts):
    return (ai_texts or {}).get(f"model_equation_{analysis_type}", "")


def _render_regresi(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Analisis Regresi Linier", level=1)
    dep_var  = data.get("y","Y")
    ind_vars = data.get("x",[])
    _add_para(doc, pr,
        f"Analisis regresi linier dilakukan untuk menguji pengaruh "
        f"{', '.join(ind_vars)} terhadap {dep_var}.",
        first_indent=True, is_last_in_block=True)

    r2, adj_r2, f_pvalue, rmse = (data.get(k) for k in ("r2","adj_r2","f_pvalue","rmse"))
    if any(v is not None for v in [r2, adj_r2, f_pvalue]):
        fd = {}
        if r2       is not None: fd["R\u00b2"]          = round(r2, 4)
        if adj_r2   is not None: fd["R\u00b2 Adjusted"] = round(adj_r2, 4)
        if f_pvalue is not None: fd["F p-value"]        = round(f_pvalue, 4)
        if rmse     is not None: fd["RMSE"]              = round(rmse, 4)
        _style_table(doc, pd.DataFrame([fd]), pr, f"Tabel {bab}.1. Ringkasan Model Regresi")

    coef_df = data.get("coef_table")
    if coef_df is not None and not coef_df.empty:
        _style_table(doc, coef_df, pr, f"Tabel {bab}.2. Koefisien Regresi")

    for ks, ls in (("_scatter","Scatter Plot Aktual vs Prediksi"),
                   ("_residual","Plot Residual"),("_koefisien","Koefisien Regresi")):
        if figs_png.get(f"regresi{ks}"):
            _embed_image(doc, pr, figs_png[f"regresi{ks}"], f"Gambar {fig_no}. {ls}")
            fig_no += 1

    if _get_eq("regresi", data, ai_texts):
        _add_model_equation_box(doc, pr, _get_eq("regresi", data, ai_texts))

    _add_ai_narasi(doc, pr,
                   (ai_texts or {}).get("regresi") or (ai_texts or {}).get("ols_plus",""),
                   "Interpretasi Regresi Linier", mod_key="regresi", data=data)
    return fig_no


def _render_ols_plus(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Regresi OLS (Diagnostik Lengkap)", level=1)
    _add_para(doc, pr,
        f"Regresi OLS dengan uji asumsi klasik untuk model "
        f"{data.get('y','Y')} ~ {', '.join(data.get('x',[]))}.",
        first_indent=True, is_last_in_block=True)

    coef_df = data.get("coef_table")
    if coef_df is not None and not coef_df.empty:
        _style_table(doc, coef_df, pr, f"Tabel {bab}.1. Koefisien Regresi OLS")

    diag = {k: round(v,4) for k,v in {
        "Durbin-Watson":  data.get("durbin_watson"),
        "VIF Maksimum":   data.get("vif_max"),
        "White Test p":   data.get("white_pvalue"),
        "Shapiro Res. p": data.get("shapiro_residual_p"),
    }.items() if v is not None}
    if diag:
        _style_table(doc, pd.DataFrame([diag]), pr, f"Tabel {bab}.2. Statistik Uji Asumsi Klasik")

    for ks, ls in (("_scatter","Scatter Aktual vs Prediksi"),
                   ("_residual","Plot Residual OLS"),("_koefisien","Koefisien OLS")):
        if figs_png.get(f"ols_plus{ks}"):
            _embed_image(doc, pr, figs_png[f"ols_plus{ks}"], f"Gambar {fig_no}. {ls}")
            fig_no += 1

    if _get_eq("ols_plus", data, ai_texts):
        _add_model_equation_box(doc, pr, _get_eq("ols_plus", data, ai_texts))

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("ols_plus",""),
                   "Interpretasi OLS", mod_key="ols_plus", data=data)
    return fig_no


def _render_logistik(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Regresi Logistik", level=1)
    _add_para(doc, pr,
        f"Regresi logistik biner untuk memprediksi {data.get('y','Y')} "
        f"berdasarkan {', '.join(data.get('x',[]))}.",
        first_indent=True, is_last_in_block=True)

    fit = {k: round(v,4) for k,v in {
        "AUC": data.get("auc"), "Pseudo R\u00b2": data.get("pseudo_r2"),
        "AIC": data.get("aic"), "BIC": data.get("bic"),
    }.items() if v is not None}
    if fit:
        _style_table(doc, pd.DataFrame([fit]), pr, f"Tabel {bab}.1. Kebaikan Model")

    coef_df = _first_valid_df(data.get("coef_table"), data.get("odds_df"))
    if coef_df is not None and not coef_df.empty:
        _style_table(doc, coef_df, pr, f"Tabel {bab}.2. Koefisien dan Odds Ratio")

    cr = data.get("cr", {})
    if cr:
        try:
            cr_df = pd.DataFrame({
                "Kelas":     ["0","1","Macro avg","Weighted avg"],
                "Precision": [round(cr[k]["precision"],4) for k in ("0","1","macro avg","weighted avg")],
                "Recall":    [round(cr[k]["recall"],4)    for k in ("0","1","macro avg","weighted avg")],
                "F1":        [round(cr[k]["f1-score"],4)  for k in ("0","1","macro avg","weighted avg")],
            })
            _style_table(doc, cr_df, pr, f"Tabel {bab}.3. Laporan Klasifikasi")
        except Exception:
            pass

    for fk, ls in (("logistik_odds_ratio","Odds Ratio"),("logistik_roc","Kurva ROC")):
        if figs_png.get(fk):
            _embed_image(doc, pr, figs_png[fk], f"Gambar {fig_no}. {ls}"); fig_no += 1

    if _get_eq("logistik", data, ai_texts):
        _add_model_equation_box(doc, pr, _get_eq("logistik", data, ai_texts))

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("logistik",""),
                   "Interpretasi Regresi Logistik", mod_key="logistik", data=data)
    return fig_no


def _render_mediasi(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Analisis Mediasi", level=1)
    x, m, y = data.get("x","X"), data.get("m","M"), data.get("y","Y")
    _add_para(doc, pr,
        f"Analisis mediasi menguji peran {m} sebagai mediator hubungan {x} \u2192 {y} "
        f"menggunakan prosedur Bootstrap (Preacher & Hayes, 2008).",
        first_indent=True, is_last_in_block=True)

    path_df = data.get("path_table")
    if path_df is not None and not path_df.empty:
        _style_table(doc, path_df, pr, f"Tabel {bab}.1. Koefisien Jalur Mediasi")

    boot_ci = data.get("bootstrap_ci")
    if boot_ci:
        ci_d = {"CI Bawah (95%)": round(boot_ci[0],4), "CI Atas (95%)": round(boot_ci[1],4)}
        if data.get("indirect_effect") is not None:
            ci_d["Efek Tidak Langsung"] = round(data["indirect_effect"],4)
        _style_table(doc, pd.DataFrame([ci_d]), pr, f"Tabel {bab}.2. Bootstrap CI")

    # Fix 4: mediasi_diagram bisa berupa bar chart efek (SVG path diagram tidak bisa jadi PNG)
    if figs_png.get("mediasi_diagram"):
        _embed_image(doc, pr, figs_png["mediasi_diagram"],
                     f"Gambar {fig_no}. Diagram Jalur Mediasi"); fig_no += 1
    elif figs_png.get("mediasi_efek"):
        _embed_image(doc, pr, figs_png["mediasi_efek"],
                     f"Gambar {fig_no}. Ringkasan Efek Mediasi (Indirect, Direct, Total)"); fig_no += 1

    if _get_eq("mediasi", data, ai_texts):
        _add_model_equation_box(doc, pr, _get_eq("mediasi", data, ai_texts))

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("mediasi",""),
                   "Interpretasi Mediasi", mod_key="mediasi", data=data)
    return fig_no


def _render_moderasi(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Analisis Moderasi", level=1)
    x, z, y = data.get("x","X"), data.get("z","Z"), data.get("y","Y")
    _add_para(doc, pr,
        f"Analisis moderasi menguji apakah {z} memoderasi hubungan {x} \u2192 {y}.",
        first_indent=True, is_last_in_block=True)

    coef_df = data.get("coef_table")
    if coef_df is not None and not coef_df.empty:
        _style_table(doc, coef_df, pr, f"Tabel {bab}.1. Koefisien Moderasi")

    r2, adj_r2 = data.get("r2"), data.get("adj_r2")
    if r2 is not None or adj_r2 is not None:
        fd = {}
        if r2:     fd["R\u00b2"]          = round(r2, 4)
        if adj_r2: fd["R\u00b2 Adjusted"] = round(adj_r2, 4)
        _style_table(doc, pd.DataFrame([fd]), pr, f"Tabel {bab}.2. Kebaikan Model Moderasi")

    if figs_png.get("moderasi_interaction"):
        _embed_image(doc, pr, figs_png["moderasi_interaction"],
                     f"Gambar {fig_no}. Plot Interaksi Moderasi"); fig_no += 1

    if _get_eq("moderasi", data, ai_texts):
        _add_model_equation_box(doc, pr, _get_eq("moderasi", data, ai_texts))

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("moderasi",""),
                   "Interpretasi Moderasi", mod_key="moderasi", data=data)
    return fig_no


def _render_anova(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Analisis Varians (ANOVA)", level=1)
    _add_para(doc, pr,
        "Uji ANOVA satu arah dilakukan untuk membandingkan rata-rata antar kelompok.",
        first_indent=True, is_last_in_block=True)

    if (df_ := data.get("anova_table")) is not None and not df_.empty:
        _style_table(doc, df_, pr, f"Tabel {bab}.1. Tabel ANOVA")
    if (df_ := data.get("posthoc_table")) is not None and not df_.empty:
        _style_table(doc, df_, pr, f"Tabel {bab}.2. Uji Post-Hoc")
    if data.get("eta_squared") is not None:
        _style_table(doc,
            pd.DataFrame([{"Eta Squared (\u03b7\u00b2)": round(data["eta_squared"],4)}]),
            pr, f"Tabel {bab}.3. Ukuran Efek")

    for fk, ls in (("anova_boxplot","Boxplot ANOVA per Kelompok"),
                   ("anova_bar","Bar Chart Rata-rata per Kelompok")):
        if figs_png.get(fk):
            _embed_image(doc, pr, figs_png[fk], f"Gambar {fig_no}. {ls}"); fig_no += 1

    if _get_eq("anova", data, ai_texts):
        _add_model_equation_box(doc, pr, _get_eq("anova", data, ai_texts))

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("anova",""),
                   "Interpretasi ANOVA", mod_key="anova", data=data)
    return fig_no


def _render_uji_beda(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Uji Beda", level=1)
    g1, g2 = data.get("g1_name","Kelompok 1"), data.get("g2_name","Kelompok 2")
    _add_para(doc, pr,
        f"Uji beda ({data.get('uji_type','t-test')}) dilakukan untuk membandingkan {g1} dan {g2}.",
        first_indent=True, is_last_in_block=True)

    res = {label: round(float(data[k]),4) for k,label in [
        ("statistic","Statistik Uji"),("p_value","p-value"),
        ("effect_size","Effect Size"),("g1_mean",f"Rata-rata {g1}"),
        ("g2_mean",f"Rata-rata {g2}"),
    ] if data.get(k) is not None}
    if res:
        _style_table(doc, pd.DataFrame([res]), pr, f"Tabel {bab}.1. Hasil Uji Beda")

    # Fix 5: embed boxplot jika tersedia
    for fk, ls in (("uji_beda_boxplot", "Boxplot Distribusi Per Kelompok"),
                   ("uji_beda_violin",  "Violin Plot Distribusi Per Kelompok")):
        if figs_png.get(fk):
            _embed_image(doc, pr, figs_png[fk], f"Gambar {fig_no}. {ls}")
            fig_no += 1

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("uji_beda",""),
                   "Interpretasi Uji Beda", mod_key="uji_beda", data=data)
    return fig_no


def _render_outlier(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Deteksi Outlier", level=1)
    if (df_ := data.get("outlier_table")) is not None and not df_.empty:
        _style_table(doc, df_, pr, f"Tabel {bab}.1. Data Outlier Terdeteksi")
    if figs_png.get("outlier_plot"):
        _embed_image(doc, pr, figs_png["outlier_plot"],
                     f"Gambar {fig_no}. Visualisasi Outlier"); fig_no += 1
    _add_ai_narasi(doc, pr, (ai_texts or {}).get("outlier",""),
                   "Interpretasi Outlier", mod_key="outlier", data=data)
    return fig_no


def _render_sem(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Structural Equation Modeling (SEM) & CFA", level=1)
    _add_para(doc, pr,
        "Analisis SEM dan CFA dilakukan untuk menguji model pengukuran dan struktural "
        "(Hair et al., 2010).",
        first_indent=True, is_last_in_block=True)

    if (df_ := data.get("fit_indices")) is not None and not df_.empty:
        _style_table(doc, df_, pr, f"Tabel {bab}.1. Indeks Kecocokan Model")
        acuan = pd.DataFrame({
            "Indeks":    ["CFI","TLI","RMSEA","SRMR","\u03c7\u00b2/df"],
            "Kriteria":  ["\u2265 .90","\u2265 .90","\u2264 .08","\u2264 .08","\u2264 2.00"],
            "Referensi": ["Hair et al. (2010)"]*5,
        })
        _style_table(doc, acuan, pr, f"Tabel {bab}.1a. Kriteria Kecocokan Model")

    if (df_ := data.get("loadings")) is not None and not df_.empty:
        _style_table(doc, df_, pr, f"Tabel {bab}.2. Factor Loadings (CFA)")
    if (df_ := data.get("path_estimates")) is not None and not df_.empty:
        _style_table(doc, df_, pr, f"Tabel {bab}.3. Estimasi Jalur Struktural")

    if figs_png.get("sem_diagram"):
        _embed_image(doc, pr, figs_png["sem_diagram"],
                     f"Gambar {fig_no}. Diagram Jalur SEM"); fig_no += 1

    if _get_eq("sem", data, ai_texts):
        _add_model_equation_box(doc, pr, _get_eq("sem", data, ai_texts))

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("sem",""),
                   "Interpretasi SEM & CFA", mod_key="sem", data=data)
    return fig_no


def _render_kelompok(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    _add_heading(doc, pr, f"{bab}. Analisis Kelompok", level=1)
    cat, num = data.get("cat","Kelompok"), data.get("num","Variabel Numerik")
    _add_para(doc, pr,
        f"Analisis perbandingan variabel '{num}' berdasarkan kelompok '{cat}'.",
        first_indent=True, is_last_in_block=True)

    if (df_ := data.get("group_stats")) is not None and not df_.empty:
        _style_table(doc, df_, pr, f"Tabel {bab}.1. Statistik per Kelompok")

    res = {label: round(float(data[k]),4) for k,label in [
        ("f_stat","F-statistik"),("p_value","p-value"),
    ] if data.get(k) is not None}
    if res:
        _style_table(doc, pd.DataFrame([res]), pr, f"Tabel {bab}.2. Hasil ANOVA Kelompok")

    if figs_png.get("kelompok_plot"):
        _embed_image(doc, pr, figs_png["kelompok_plot"],
                     f"Gambar {fig_no}. Visualisasi Kelompok"); fig_no += 1

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("kelompok",""),
                   "Interpretasi Kelompok", mod_key="kelompok", data=data)
    return fig_no



def _render_efa(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Fix 9: EFA renderer — sebelumnya tidak ada di _MODULE_RENDERERS."""
    _add_heading(doc, pr, f"{bab}. Analisis Faktor Eksploratori (EFA)", level=1)
    _add_para(doc, pr,
        "Analisis Faktor Eksploratori (EFA) dilakukan untuk mengidentifikasi struktur "
        "laten yang mendasari item-item pengukuran (Hair et al., 2010).",
        first_indent=True, is_last_in_block=True)

    # KMO & Bartlett
    kmo    = data.get("kmo")
    kmo_lbl= data.get("kmo_label", "")
    bart_p = data.get("bartlett_p")
    if kmo is not None or bart_p is not None:
        fit_d = {}
        if kmo     is not None: fit_d["KMO"]                       = round(float(kmo), 4)
        if kmo_lbl:             fit_d["Kategori KMO"]              = kmo_lbl
        if bart_p  is not None: fit_d["Bartlett p-value"]          = round(float(bart_p), 4)
        n_fac = data.get("n_factors")
        rot   = data.get("rotation", "")
        tot_v = data.get("total_var")
        if n_fac  is not None: fit_d["Jumlah Faktor"]              = int(n_fac)
        if rot:                fit_d["Rotasi"]                      = rot
        if tot_v  is not None: fit_d["Total Variance Explained (%)"] = round(float(tot_v), 2)
        _style_table(doc, pd.DataFrame([fit_d]), pr,
                     f"Tabel {bab}.1. Hasil Uji Kelayakan Data EFA")

    # Variance explained per factor
    var_df = data.get("variance_df")
    if var_df is not None and not var_df.empty:
        _style_table(doc, var_df, pr, f"Tabel {bab}.2. Variansi yang Dijelaskan per Faktor")

    # Factor loadings
    load_df = data.get("loading_df")
    if load_df is not None and not load_df.empty:
        _style_table(doc, load_df, pr, f"Tabel {bab}.3. Factor Loadings")

    # Scree plot / diagram
    for fk, ls in (("efa_scree",   "Scree Plot"),
                   ("efa_loading", "Heatmap Factor Loadings")):
        if figs_png.get(fk):
            _embed_image(doc, pr, figs_png[fk], f"Gambar {fig_no}. {ls}")
            fig_no += 1

    # AI narasi — ambil dari ai_text field dalam data jika ada
    narasi = (ai_texts or {}).get("efa", "") or data.get("ai_text", "")
    _add_ai_narasi(doc, pr, narasi, "Interpretasi EFA", mod_key="efa", data=data)
    return fig_no

# =============================================================================
# RENDERER BARU — Poin 3: ols_robust, ols_wls, ols_robust_comparison,
#                          compute, reliabilitas_icc, uji_asumsi
# =============================================================================

def _render_ols_robust(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk Regresi Robust (RLM Huber-M / Bisquare)."""
    _add_heading(doc, pr, f"{bab}. Regresi Robust (RLM)", level=1)
    dep_var   = data.get("dep_var", "Y")
    ind_vars  = data.get("ind_vars", [])
    estimator = data.get("estimator", "Huber-M")
    n_obs     = data.get("n_obs", "?")
    n_low     = data.get("n_low_weight", 0)
    n_changed = data.get("n_changed", 0)

    _add_para(doc, pr,
        f"Regresi Robust ({estimator}) dilakukan untuk memodelkan {dep_var} "
        f"berdasarkan {', '.join(ind_vars) if ind_vars else 'variabel prediktor'} "
        f"(N = {n_obs}). Estimator robust mereduksi pengaruh outlier dan leverage "
        f"points dengan memberi bobot rendah pada observasi bermasalah "
        f"(Huber, 1973; Greene, 2012).",
        first_indent=True, is_last_in_block=True)

    # Ringkasan metrik
    summary = {}
    if n_obs:          summary["N Observasi"]           = n_obs
    if estimator:      summary["Estimator"]              = estimator
    if n_low is not None: summary["Obs. Downweighted (<0.5)"] = n_low
    if n_changed is not None: summary["Koef. Berubah >10%"]   = n_changed
    ols_rmse = data.get("ols_rmse")
    rlm_rmse = data.get("rlm_rmse")
    if ols_rmse is not None: summary["RMSE OLS"]  = round(float(ols_rmse), 4)
    if rlm_rmse is not None: summary["RMSE RLM"]  = round(float(rlm_rmse), 4)
    if summary:
        _style_table(doc, pd.DataFrame([summary]), pr,
                     f"Tabel {bab}.1. Ringkasan Model Regresi Robust")

    # Tabel perbandingan koefisien OLS vs RLM
    coef_df = data.get("coef_df")
    if coef_df is not None and not coef_df.empty:
        _style_table(doc, coef_df, pr,
                     f"Tabel {bab}.2. Perbandingan Koefisien OLS vs {estimator}")

    # Figures
    for fk, ls in (("robust_weights",   "Robust Weights vs Residual"),
                   ("robust_coef_comp", "Perbandingan Koefisien OLS vs RLM")):
        if figs_png.get(fk):
            _embed_image(doc, pr, figs_png[fk], f"Gambar {fig_no}. {ls}")
            fig_no += 1

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("ols_robust", ""),
                   "Interpretasi Regresi Robust", mod_key="ols_robust", data=data)
    return fig_no


def _render_ols_wls(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk Weighted Least Squares (WLS)."""
    _add_heading(doc, pr, f"{bab}. Weighted Least Squares (WLS)", level=1)
    dep_var       = data.get("dep_var", "Y")
    ind_vars      = data.get("ind_vars", [])
    weight_method = data.get("weight_method", "1/|ε|")
    n_obs         = data.get("n_obs", "?")

    _add_para(doc, pr,
        f"WLS dilakukan untuk mengatasi heteroskedastisitas pada model "
        f"{dep_var} ~ {', '.join(ind_vars) if ind_vars else 'prediktor'} "
        f"menggunakan pembobotan {weight_method} (Greene, 2012).",
        first_indent=True, is_last_in_block=True)

    # Perbandingan metrik
    summary = {}
    r2      = data.get("r2")
    adj_r2  = data.get("adj_r2")
    ols_g   = data.get("ols_glejser_p")
    wls_g   = data.get("wls_glejser_p")
    ols_r   = data.get("ols_rmse")
    wls_r   = data.get("wls_rmse")
    n_ch    = data.get("n_changed")
    if r2      is not None: summary["R² WLS"]             = round(float(r2),    4)
    if adj_r2  is not None: summary["R² Adj WLS"]         = round(float(adj_r2), 4)
    if ols_g   is not None: summary["Glejser p (OLS)"]    = round(float(ols_g),  4)
    if wls_g   is not None: summary["Glejser p (WLS)"]    = round(float(wls_g),  4)
    if ols_r   is not None: summary["RMSE OLS"]           = round(float(ols_r),  4)
    if wls_r   is not None: summary["RMSE WLS"]           = round(float(wls_r),  4)
    if n_ch    is not None: summary["Koef. Berubah >10%"] = n_ch
    if summary:
        _style_table(doc, pd.DataFrame([summary]), pr,
                     f"Tabel {bab}.1. Perbandingan OLS vs WLS")

    # Tabel koefisien
    coef_df = data.get("coef_df")
    if coef_df is not None and not coef_df.empty:
        _style_table(doc, coef_df, pr,
                     f"Tabel {bab}.2. Koefisien OLS vs WLS")

    # Figures
    for fk, ls in (("wls_fit_plot",   "Plot Aktual vs Prediksi: OLS vs WLS"),
                   ("wls_coef_comp",  "Perbandingan Koefisien OLS vs WLS")):
        if figs_png.get(fk):
            _embed_image(doc, pr, figs_png[fk], f"Gambar {fig_no}. {ls}")
            fig_no += 1

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("ols_wls", ""),
                   "Interpretasi WLS", mod_key="ols_wls", data=data)
    return fig_no


def _render_ols_robust_comparison(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk Perbandingan Model: OLS / RLM-Huber / RLM-Bisquare / WLS."""
    _add_heading(doc, pr, f"{bab}. Perbandingan Model Regresi", level=1)
    dep_var    = data.get("dep_var", "Y")
    ind_vars   = data.get("ind_vars", [])
    best_model = data.get("best_model", "OLS (Baseline)")
    n_obs      = data.get("n_obs", "?")

    _add_para(doc, pr,
        f"Perbandingan empat model regresi (OLS, RLM Huber-M, RLM Bisquare, WLS) "
        f"dilakukan untuk memilih pendekatan estimasi terbaik bagi model "
        f"{dep_var} ~ {', '.join(ind_vars) if ind_vars else 'prediktor'} "
        f"(N = {n_obs}). Model terbaik dipilih berdasarkan RMSE terendah.",
        first_indent=True, is_last_in_block=True)

    comp_df = data.get("comparison_df")
    if comp_df is not None and not comp_df.empty:
        _style_table(doc, comp_df, pr,
                     f"Tabel {bab}.1. Perbandingan Metrik Semua Model")

    _add_para(doc, pr,
        f"Berdasarkan evaluasi RMSE, model terbaik adalah: {best_model}.",
        first_indent=True, is_last_in_block=True)

    if figs_png.get("robust_comparison_bar"):
        _embed_image(doc, pr, figs_png["robust_comparison_bar"],
                     f"Gambar {fig_no}. RMSE per Model")
        fig_no += 1

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("ols_robust_comparison", ""),
                   "Rekomendasi Pemilihan Model",
                   mod_key="ols_robust_comparison", data=data)
    return fig_no


def _render_compute(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk Compute Variabel Baru."""
    _add_heading(doc, pr, f"{bab}. Transformasi & Komputasi Variabel", level=1)

    compute_log = data.get("compute_log", [])
    n_ops = len(compute_log)

    _add_para(doc, pr,
        f"Sejumlah {n_ops} variabel baru dibuat melalui operasi komputasi "
        f"(formula kustom, skor komposit, recode, standardisasi, atau transformasi) "
        f"sebelum analisis utama dilakukan.",
        first_indent=True, is_last_in_block=True)

    if compute_log:
        log_rows = []
        for entry in compute_log:
            log_rows.append({
                "Variabel Baru": entry.get("new_col", ""),
                "Metode":        entry.get("method", ""),
                "Sumber/Formula": str(entry.get("source", ""))[:80],
            })
        if log_rows:
            _style_table(doc, pd.DataFrame(log_rows), pr,
                         f"Tabel {bab}.1. Log Operasi Compute Variabel")

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("compute", ""),
                   "Rasionalisasi Komputasi Variabel",
                   mod_key="compute", data=data)
    return fig_no


def _render_reliabilitas_icc(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk Uji Reliabilitas ICC."""
    _add_heading(doc, pr, f"{bab}. Uji Reliabilitas ICC (Intraclass Correlation)", level=1)

    n_subj      = data.get("n_subj", "?")
    n_rater     = data.get("n_rater", "?")
    use_type    = data.get("use_type", "")
    rec_model   = data.get("rec_model", "ICC(2,1)")
    rater_names = data.get("rater_names", [])

    _add_para(doc, pr,
        f"Uji reliabilitas menggunakan Intraclass Correlation Coefficient (ICC) "
        f"dilakukan untuk mengevaluasi konsistensi pengukuran "
        f"({use_type if use_type else 'rater agreement'}) "
        f"dengan {n_subj} subjek dan {n_rater} rater/sesi pengukuran "
        f"(Koo & Mae, 2016; Shrout & Fleiss, 1979).",
        first_indent=True, is_last_in_block=True)

    # Tabel hasil ICC semua model
    icc_records = data.get("icc_df", [])
    if icc_records:
        try:
            icc_df = pd.DataFrame(icc_records)
            # Pilih kolom utama saja agar tabel tidak terlalu lebar
            cols_show = [c for c in ("Model", "Tipe", "ICC", "CI_Lower",
                                     "CI_Upper", "F", "p_value")
                         if c in icc_df.columns]
            if cols_show:
                icc_show = icc_df[cols_show].copy()
                icc_show.columns = [c.replace("_", " ") for c in cols_show]
                _style_table(doc, icc_show, pr,
                             f"Tabel {bab}.1. Hasil Uji Reliabilitas ICC — Semua Model")
        except Exception:
            pass

    # Panduan interpretasi threshold
    threshold_df = pd.DataFrame({
        "Rentang ICC":        ["ICC < 0.50", "0.50 ≤ ICC < 0.75",
                               "0.75 ≤ ICC < 0.90", "ICC ≥ 0.90"],
        "Kualitas":           ["Buruk", "Sedang", "Baik", "Sangat Baik (Excellent)"],
        "Sumber":             ["Koo & Mae (2016)"] * 4,
    })
    _style_table(doc, threshold_df, pr,
                 f"Tabel {bab}.2. Panduan Interpretasi ICC (Koo & Mae, 2016)")

    # Tabel ANOVA dasar perhitungan
    anova_records = data.get("anova_tbl", [])
    if anova_records:
        try:
            anova_df = pd.DataFrame(anova_records)
            if not anova_df.empty:
                _style_table(doc, anova_df, pr,
                             f"Tabel {bab}.3. Tabel ANOVA Dasar Perhitungan ICC")
        except Exception:
            pass

    # AI narasi
    ai_text = (ai_texts or {}).get("reliabilitas_icc", "") or data.get("ai_text", "")
    _add_ai_narasi(doc, pr, ai_text,
                   "Interpretasi Reliabilitas ICC",
                   mod_key="reliabilitas_icc", data=data)
    return fig_no


def _render_uji_asumsi(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk Uji Asumsi Pra-Analisis."""
    _add_heading(doc, pr, f"{bab}. Uji Asumsi Pra-Analisis", level=1)
    _add_para(doc, pr,
        "Uji asumsi dilakukan sebelum pemilihan metode analisis utama, mencakup "
        "normalitas multivariat (Mardia's test), homogenitas varians (Levene & Bartlett), "
        "dan linieritas (Ramsey RESET test).",
        first_indent=True, is_last_in_block=True)

    # Rekomendasi
    rec = data.get("rekomendasi", {})
    if isinstance(rec, dict):
        level      = rec.get("level", "").upper()
        skor       = rec.get("skor_lulus", "?")
        total      = rec.get("total_uji", "?")
        pct        = rec.get("pct_lulus", "?")
        rec_list   = rec.get("rekomendasi", [])
        peringatan = rec.get("peringatan", [])

        _style_table(doc, pd.DataFrame([{
            "Level Analisis": level,
            "Uji Lulus":      f"{skor}/{total}",
            "% Lulus":        f"{pct}%",
        }]), pr, f"Tabel {bab}.1. Ringkasan Hasil Uji Asumsi")

        if rec_list:
            _add_heading(doc, pr, "Rekomendasi Metode Analisis", level=3)
            clean_recs = [r.lstrip("✅🟡🔴 ").replace("**", "") for r in rec_list]
            for r in clean_recs:
                _add_para(doc, pr, f"• {r}", first_indent=False)

        if peringatan:
            _add_heading(doc, pr, "Peringatan", level=3)
            clean_warns = [w.lstrip("⚠️🚨ℹ️ ").replace("**", "") for w in peringatan]
            for w in clean_warns:
                _add_para(doc, pr, f"• {w}", first_indent=False)

    # Tabel detail
    detail = rec.get("detail", {}) if isinstance(rec, dict) else {}
    if detail:
        detail_rows = []
        labels_map = {
            "normalitas_univariat":   "Normalitas Univariat (Shapiro-Wilk)",
            "normalitas_multivariat": "Normalitas Multivariat (Mardia)",
            "homogenitas":            "Homogenitas Varians (Levene)",
            "linieritas":             "Linieritas (Ramsey RESET)",
        }
        for key, info in detail.items():
            if not isinstance(info, dict):
                continue
            lulus = info.get("lulus", False)
            label = labels_map.get(key, key)
            # Buat keterangan ringkas
            if key == "normalitas_univariat":
                ket = f"{info.get('n_normal','?')}/{info.get('n_variabel','?')} variabel normal ({info.get('pct_normal','?')}%)"
            elif key == "normalitas_multivariat":
                ket = f"Skewness p={info.get('p_skew','?')}, Kurtosis p={info.get('p_kurt','?')}"
            elif key == "homogenitas":
                ket = f"{info.get('n_lulus','?')}/{info.get('n_uji','?')} kelompok homogen"
            elif key == "linieritas":
                ket = f"{info.get('n_lulus','?')}/{info.get('n_uji','?')} pasangan linier"
            else:
                ket = ""
            detail_rows.append({
                "Uji Asumsi": label,
                "Status":     "Terpenuhi ✓" if lulus else "Tidak Terpenuhi ✗",
                "Keterangan": ket,
            })
        if detail_rows:
            _style_table(doc, pd.DataFrame(detail_rows), pr,
                         f"Tabel {bab}.2. Detail Hasil Per Uji Asumsi")

    ai_text = (ai_texts or {}).get("uji_asumsi", "")
    _add_ai_narasi(doc, pr, ai_text,
                   "Interpretasi Uji Asumsi",
                   mod_key="uji_asumsi", data=data)
    return fig_no



def _render_klaster(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk Analisis Klaster (K-Means / Hierarchical)."""
    _add_heading(doc, pr, f"{bab}. Analisis Klaster", level=1)
    method    = data.get("method", "K-Means")
    k         = data.get("k", "?")
    cols      = data.get("cols", [])
    silhouette= data.get("silhouette")
    linkage   = data.get("linkage", "ward")
    n_var     = len(cols)

    sil_kat = ""
    if silhouette is not None:
        sv = float(silhouette)
        sil_kat = ("sangat baik" if sv >= 0.70 else
                   "baik" if sv >= 0.50 else
                   "cukup" if sv >= 0.30 else "rendah")

    _add_para(doc, pr,
        f"Analisis klaster dilakukan menggunakan metode {method} "
        f"dengan {k} klaster pada {n_var} variabel: "
        f"{', '.join(cols) if cols else 'lihat konfigurasi'}. "
        f"Data distandarisasi (Z-score) sebelum clustering untuk menghilangkan "
        f"efek skala (MacQueen, 1967; Ward, 1963).",
        first_indent=True, is_last_in_block=True)

    # Ringkasan metrik
    summary = {"Metode": method, "k Klaster": k, "N Variabel": n_var}
    if silhouette is not None:
        summary["Silhouette Score"] = round(float(silhouette), 4)
        summary["Kualitas"]         = sil_kat
    if method == "Hierarchical":
        summary["Linkage"] = linkage
    _style_table(doc, pd.DataFrame([summary]), pr,
                 f"Tabel {bab}.1. Ringkasan Hasil Analisis Klaster")

    # Tabel profil klaster
    profile_records = data.get("profile_df_records")
    if profile_records:
        try:
            profile_df = pd.DataFrame(profile_records)
            # Pilih kolom yang tidak terlalu lebar (max 8)
            show_cols = profile_df.columns.tolist()
            if len(show_cols) > 9:
                mean_cols = [c for c in show_cols if "(mean)" in c][:6]
                show_cols = ["Klaster", "N"] + mean_cols
                show_cols = [c for c in show_cols if c in profile_df.columns]
            _style_table(doc, profile_df[show_cols], pr,
                         f"Tabel {bab}.2. Profil Rata-rata per Klaster")
        except Exception:
            pass

    # Interpretasi silhouette standar
    if silhouette is not None:
        _add_para(doc, pr,
            f"Silhouette score = {float(silhouette):.4f} mengindikasikan kualitas "
            f"pengelompokan yang {sil_kat} (Rousseeuw, 1987). "
            f"Nilai ini menunjukkan bahwa klaster yang terbentuk memiliki "
            f"kohesi internal yang {'memadai' if float(silhouette) >= 0.30 else 'perlu diperbaiki'}.",
            first_indent=True, is_last_in_block=True)

    # Figures
    for fk, ls in (("klaster_scatter",  "Scatter Klaster via PCA"),
                   ("klaster_radar",    "Radar Chart Profil Klaster"),
                   ("klaster_elbow",    "Elbow Method & Silhouette Score"),
                   ("klaster_dendro",   "Dendrogram Hierarkikal")):
        if figs_png.get(fk):
            _embed_image(doc, pr, figs_png[fk], f"Gambar {fig_no}. {ls}")
            fig_no += 1

    _add_ai_narasi(doc, pr, (ai_texts or {}).get("klaster", ""),
                   "Interpretasi Analisis Klaster",
                   mod_key="klaster", data=data)
    return fig_no


def _render_eda(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk EDA — ringkasan profil dataset."""
    _add_heading(doc, pr, f"{bab}. Eksplorasi Data (EDA)", level=1)
    n_rows    = data.get("n_rows", "?")
    n_cols    = data.get("n_cols", "?")
    n_num     = data.get("n_numeric", "?")
    n_cat     = data.get("n_cat", "?")
    n_miss    = data.get("n_missing", 0)
    pct_miss  = data.get("pct_missing", 0)
    n_dup     = data.get("n_dup", 0)

    _add_para(doc, pr,
        f"Eksplorasi data awal (EDA) dilakukan terhadap dataset yang terdiri dari "
        f"{n_rows} observasi dan {n_cols} variabel ({n_num} numerik, {n_cat} kategorik). "
        f"Terdapat {n_miss} missing values ({pct_miss}% dari total sel) "
        f"dan {n_dup} baris duplikat.",
        first_indent=True, is_last_in_block=True)

    # Tabel ringkasan profil
    summary = {
        "Jumlah Observasi": n_rows,
        "Jumlah Variabel":  n_cols,
        "Variabel Numerik": n_num,
        "Variabel Kategorik": n_cat,
        "Total Missing":    n_miss,
        "% Missing":        f"{pct_miss}%",
        "Baris Duplikat":   n_dup,
    }
    _style_table(doc, pd.DataFrame([summary]), pr,
                 f"Tabel {bab}.1. Ringkasan Profil Dataset")

    # Interpretasi kondisi data
    kondisi = []
    if float(pct_miss) > 20:
        kondisi.append(f"missing values tinggi ({pct_miss}%) — perlu imputasi atau eliminasi kolom")
    elif float(pct_miss) > 5:
        kondisi.append(f"missing values moderat ({pct_miss}%) — pertimbangkan imputasi")
    if n_dup > 0:
        kondisi.append(f"terdapat {n_dup} baris duplikat yang perlu diperiksa")

    if kondisi:
        _add_para(doc, pr,
            "Perhatian pada kualitas data: " + "; ".join(kondisi) + ".",
            first_indent=True, is_last_in_block=True)
    else:
        _add_para(doc, pr,
            "Dataset dalam kondisi baik untuk dilanjutkan ke tahap analisis statistik.",
            first_indent=True, is_last_in_block=True)

    # Figures EDA jika ada
    for fk, ls in (("eda_profil",    "Komposisi Tipe Kolom"),
                   ("eda_missing",   "Missing Values per Kolom"),
                   ("eda_histogram", "Distribusi Variabel Numerik")):
        if figs_png.get(fk):
            _embed_image(doc, pr, figs_png[fk], f"Gambar {fig_no}. {ls}")
            fig_no += 1

    return fig_no


def _render_cfa(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk CFA Standalone — Confirmatory Factor Analysis."""
    _add_heading(doc, pr, f"{bab}. Confirmatory Factor Analysis (CFA)", level=1)

    factor_map  = data.get("factor_map", {})
    model_syntax= data.get("model_syntax", "")
    n_obs       = data.get("n_obs", "?")
    n_konstruk  = len(factor_map)

    _add_para(doc, pr,
        f"Confirmatory Factor Analysis (CFA) dilakukan untuk menguji validitas "
        f"konstruk model pengukuran dengan {n_konstruk} konstruk laten "
        f"dan {n_obs} observasi. "
        f"Estimasi menggunakan metode Maximum Likelihood via semopy "
        f"(Hair et al., 2010; Hu & Bentler, 1999).",
        first_indent=True, is_last_in_block=True)

    if model_syntax:
        _add_heading(doc, pr, "Sintaks Model", level=3)
        _add_para(doc, pr, model_syntax, first_indent=False, is_last_in_block=True)

    # Tabel fit indices
    fit_records = data.get("fit_df_records")
    if fit_records:
        try:
            _style_table(doc, pd.DataFrame(fit_records), pr,
                         f"Tabel {bab}.1. Indeks Kecocokan Model (Goodness of Fit)")
        except Exception:
            pass

    # Tabel factor loadings
    loading_records = data.get("loadings_df_records")
    if loading_records:
        try:
            ldf = pd.DataFrame(loading_records)
            show_cols = [c for c in ("Konstruk Laten", "Indikator", "Loading (λ)",
                                     "SE", "z / t", "p-value", "Kecukupan")
                         if c in ldf.columns]
            _style_table(doc, ldf[show_cols] if show_cols else ldf, pr,
                         f"Tabel {bab}.2. Factor Loadings (λ)")
        except Exception:
            pass

    # Tabel AVE & CR
    ave_records = data.get("ave_cr_df_records")
    if ave_records:
        try:
            _style_table(doc, pd.DataFrame(ave_records), pr,
                         f"Tabel {bab}.3. AVE dan Composite Reliability (CR)")
        except Exception:
            pass

    # Tabel HTMT
    htmt_records = data.get("htmt_df_records")
    if htmt_records:
        try:
            htmt_df = pd.DataFrame(htmt_records)
            _style_table(doc, htmt_df, pr,
                         f"Tabel {bab}.4. HTMT Matrix (Discriminant Validity)")
        except Exception:
            pass

    # Figure loadings
    if figs_png.get("cfa_loadings"):
        _embed_image(doc, pr, figs_png["cfa_loadings"],
                     f"Gambar {fig_no}. Factor Loadings per Konstruk")
        fig_no += 1

    ai_text = (ai_texts or {}).get("cfa", "")
    _add_ai_narasi(doc, pr, ai_text,
                   "Interpretasi CFA", mod_key="cfa", data=data)
    return fig_no


def _render_scraping(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Renderer untuk Web Scraping & Data Collector."""
    _add_heading(doc, pr, f"{bab}. Pengumpulan Data via Web Scraping", level=1)

    source   = data.get("source", "tidak diketahui")
    n_rows   = data.get("n_rows", "?")
    n_cols   = data.get("n_cols", "?")
    n_num    = data.get("n_numeric", "?")
    n_miss   = data.get("n_missing", 0)
    n_dup    = data.get("n_dup", 0)
    cols     = data.get("col_names", [])

    _add_para(doc, pr,
        f"Data penelitian dikumpulkan secara otomatis melalui teknik web scraping "
        f"dari sumber: {source}. "
        f"Dataset yang diperoleh terdiri dari {n_rows} observasi dan {n_cols} variabel "
        f"({n_num} numerik). "
        f"Proses pengambilan data menggunakan pustaka requests dan BeautifulSoup "
        f"dengan rate-limiting untuk mematuhi etika pengambilan data publik.",
        first_indent=True, is_last_in_block=True)

    # Tabel ringkasan
    summary = {
        "Sumber Data":      source[:80],
        "Jumlah Observasi": n_rows,
        "Jumlah Variabel":  n_cols,
        "Variabel Numerik": n_num,
        "Missing Values":   n_miss,
        "Baris Duplikat":   n_dup,
    }
    _style_table(doc, pd.DataFrame([summary]), pr,
                 f"Tabel {bab}.1. Ringkasan Dataset Hasil Scraping")

    if cols:
        cols_df = pd.DataFrame({"No": range(1, len(cols)+1), "Nama Variabel": cols})
        _style_table(doc, cols_df, pr,
                     f"Tabel {bab}.2. Daftar Variabel")

    _add_para(doc, pr,
        f"Setelah pengambilan data, dilakukan proses cleaning meliputi: "
        f"penghapusan kolom dengan missing values tinggi, standardisasi format angka, "
        f"penghapusan duplikat, dan konversi tipe data. "
        + (f"Terdapat {n_miss} missing values yang ditangani sebelum analisis. "
           if int(n_miss) > 0 else "Dataset hasil scraping tidak memiliki missing values. "),
        first_indent=True, is_last_in_block=True)

    ai_text = (ai_texts or {}).get("scraping", "")
    _add_ai_narasi(doc, pr, ai_text,
                   "Evaluasi Kualitas Data Scraping",
                   mod_key="scraping", data=data)
    return fig_no


_MODULE_RENDERERS = {
    "regresi": _render_regresi,
    "ols_plus":              _render_ols_plus,
    "logistik": _render_logistik,
    "mediasi":               _render_mediasi,
    "moderasi":              _render_moderasi,
    "anova":                 _render_anova,
    "uji_beda":              _render_uji_beda,
    "outlier":               _render_outlier,
    "sem":                   _render_sem,
    "kelompok":              _render_kelompok,
    "efa":                   _render_efa,
    # Poin 3: renderer baru
    "ols_robust":            _render_ols_robust,
    "ols_wls":               _render_ols_wls,
    "ols_robust_comparison": _render_ols_robust_comparison,
    "compute":               _render_compute,
    "reliabilitas_icc":      _render_reliabilitas_icc,
    "uji_asumsi":            _render_uji_asumsi,
    # Modul baru v4.2
    "klaster":               _render_klaster,
    "eda":                   _render_eda,
    "cfa":                   _render_cfa,
    "scraping":              _render_scraping,
    "power_analysis":       _render_power_analysis,
}


# =============================================================================
# TITLE PAGE — satu-satunya tempat blank paragraf yang memang by design
# =============================================================================

def _render_power_analysis(doc, pr, data, figs_png, fig_no, bab, ai_texts):
    """Render hasil Power Analysis ke dalam dokumen Word."""
    _add_heading(doc, pr, "Power Analysis", level=2)

    # Ringkasan parameter
    params = {
        k: v for k, v in {
            "Jenis Uji":       data.get("test_type"),
            "Effect Size":     data.get("effect_size"),
            "Alpha (α)": data.get("alpha"),
            "Power (1−β)": data.get("power"),
            "N Total":         data.get("n_total"),
            "N per Kelompok":  data.get("n_per_group"),
        }.items() if v is not None
    }
    if params:
        _style_table(
            doc, pd.DataFrame([params]), pr,
            f"Tabel {bab}.1. Parameter Power Analysis",
        )

    # Narasi
    _add_ai_narasi(
        doc, pr,
        ai_texts.get("power_analysis") or _fallback_narasi("power_analysis", data),
        f"Tabel {bab}.1",
    )


def _add_title_page(doc: Document, p: StyleProfile,
                    report, cols, report_style, data_type,
                    user_name: str = ""):
    # Gunakan space_before besar di paragraf judul, bukan blank lines
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before      = Pt(120)   # ~4 baris vertikal
    pt.paragraph_format.space_after       = Pt(0)
    pt.paragraph_format.line_spacing      = p.line_spacing
    pt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pt.paragraph_format.first_line_indent = Pt(0)
    r = pt.add_run("LAPORAN ANALISIS STATISTIK" if p.h1_upper else "Laporan Analisis Statistik")
    r.font.name = p.font; r.font.size = p.font_size
    r.bold = True; r.font.color.rgb = p.h1_color

    _cover_lines = [
        f"Format: {report_style}",
        f"Tipe Data: {data_type.split('(')[0].strip()}",
        f"Total Responden: {report.get('rows_after_clean','?')}",
        f"Variabel Analisis: {len(cols)}",
        f"Tanggal: {datetime.date.today().strftime('%d %B %Y')}",
        "Ruang Statistika \u2014 AI-Powered Research & Stats Reporting",
    ]
    if user_name:
        _cover_lines.insert(4, f"Peneliti: {user_name}")

    for i, line in enumerate(_cover_lines):
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.paragraph_format.space_before      = Pt(24) if i == 0 else Pt(0)
        pi.paragraph_format.space_after       = Pt(0)
        pi.paragraph_format.line_spacing      = p.line_spacing
        pi.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pi.paragraph_format.first_line_indent = Pt(0)
        ri = pi.add_run(line)
        ri.font.name = p.font; ri.font.size = p.font_size
    doc.add_page_break()


# =============================================================================
# REFERENCES
# =============================================================================

def _add_references_section(doc: Document, p: StyleProfile, apa_refs):
    _add_heading(doc, p, p.ref_label, level=1)
    refs = apa_refs if isinstance(apa_refs, list) else \
           [l for l in str(apa_refs).split("\n") if l.strip()]
    for idx, ref in enumerate(refs, 1):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = para.paragraph_format
        fmt.line_spacing      = p.line_spacing
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.space_before      = Pt(0)
        fmt.space_after       = Pt(0)
        if p.ref_hanging_indent:
            fmt.first_line_indent = -Cm(1.27)
            fmt.left_indent       = Cm(1.27)
            run_text = str(ref).strip()
        else:
            fmt.first_line_indent = Pt(0)
            run_text = f"{idx}. {str(ref).strip()}"
        run = para.add_run(run_text)
        run.font.name = p.font; run.font.size = p.font_size


# =============================================================================
# MAIN GENERATOR
# =============================================================================

def generate_pro_docx(
    df,
    report: dict,
    desc_df,
    norm_df,
    val_df,
    alpha_val,
    corr_matrix,
    cols: list,
    r_tabel: float,
    figs: dict | None = None,
    figs_png: dict | None = None,
    ai_texts: dict | None = None,
    report_style: str = "APA 7th Edition",
    data_type: str = "Data Primer (Kuesioner / Survei)",
    session_results: dict | None = None,
    user_name: str = "",
) -> io.BytesIO:

    try:
        from utils.stats_helpers import narrate_descriptive, narrate_validity
        _has_narr = True
    except ImportError:
        _has_narr = False

    if figs            is None: figs            = {}
    if figs_png        is None: figs_png        = {}
    if ai_texts        is None: ai_texts        = {}
    if session_results is None: session_results = {}

    profile = _get_profile(report_style)

    for key, fig in figs.items():
        if key not in figs_png:
            try:
                figs_png[key] = fig.to_image(format="png", width=700, height=400, scale=1.5)
            except Exception:
                pass

    doc = Document()
    _set_doc_defaults(doc, profile)

    fig_ctr = 1
    bab_ctr = 1

    _add_title_page(doc, profile, report, cols, report_style, data_type, user_name=user_name)

    # BAB 1: Ringkasan Data
    _add_heading(doc, profile, f"{bab_ctr}. Ringkasan Data dan Pembersihan", level=1)
    bab_ctr += 1
    rows_orig  = report.get("original_rows",   len(df))
    cols_orig  = report.get("original_cols",   len(df.columns))
    rows_clean = report.get("rows_after_clean", len(df))
    missing    = report.get("missing_per_col",  {})

    _add_para(doc, profile,
        f"Dataset awal memiliki {rows_orig} baris dan {cols_orig} kolom. "
        f"Setelah pembersihan data, terdapat {rows_clean} baris yang valid untuk dianalisis.",
        first_indent=True, is_last_in_block=True)

    if missing:
        _add_para(doc, profile,
            "Nilai yang hilang ditemukan pada variabel: " +
            ", ".join([f"{k} ({v} nilai)" for k,v in missing.items()]) + ".",
            first_indent=True, is_last_in_block=True)
    else:
        _add_para(doc, profile,
            "Tidak ditemukan nilai yang hilang (missing values) pada dataset.",
            first_indent=True, is_last_in_block=True)

    _style_table(doc,
        pd.DataFrame({
            "Keterangan": ["Baris Awal","Baris Valid","Kolom","Variabel Analisis"],
            "Jumlah":     [rows_orig, rows_clean, cols_orig, len(cols)],
        }), profile, "Tabel 1.1. Ringkasan Dataset")

    # BAB 2: Deskriptif
    if desc_df is not None and not desc_df.empty:
        _add_heading(doc, profile, f"{bab_ctr}. Statistik Deskriptif", level=1)
        bab_no = bab_ctr; bab_ctr += 1
        _style_table(doc, desc_df, profile, f"Tabel {bab_no}.1. Statistik Deskriptif Variabel")
        narr = ai_texts.get("descriptive","")
        if not narr and _has_narr:
            try: narr = narrate_descriptive(desc_df).replace("**","")
            except Exception: narr = ""
        _add_ai_narasi(doc, profile, narr, "Interpretasi Statistik Deskriptif")
        if figs_png.get("histogram"):
            _embed_image(doc, profile, figs_png["histogram"],
                         f"Gambar {fig_ctr}. Distribusi Variabel Penelitian"); fig_ctr += 1

    # BAB 3: Normalitas
    if norm_df is not None and not norm_df.empty:
        _add_heading(doc, profile, f"{bab_ctr}. Uji Normalitas (Shapiro-Wilk)", level=1)
        bab_no = bab_ctr; bab_ctr += 1
        _add_para(doc, profile,
            "Uji Shapiro-Wilk digunakan untuk menguji asumsi normalitas. "
            "H\u2080 ditolak apabila nilai p < .05.",
            first_indent=True, is_last_in_block=True)
        _style_table(doc, norm_df, profile,
                     f"Tabel {bab_no}.1. Hasil Uji Normalitas Shapiro-Wilk")
        # Fallback narasi normalitas otomatis
        narr = ai_texts.get("normality","")
        if not narr:
            try:
                normals = norm_df[norm_df.iloc[:,2].astype(str).str.lower().isin(["ya","yes","normal","true"])].iloc[:,0].tolist() \
                          if norm_df.shape[1] >= 3 else []
                all_n   = norm_df.shape[0]
                if len(normals) == all_n:
                    narr = "Hasil uji Shapiro-Wilk menunjukkan bahwa seluruh variabel berdistribusi normal (p > .05), sehingga asumsi normalitas terpenuhi untuk analisis parametrik selanjutnya."
                elif normals:
                    narr = (f"Sebagian variabel ({', '.join(normals)}) berdistribusi normal (p > .05). "
                            f"Variabel lainnya tidak memenuhi asumsi normalitas, sehingga perlu dipertimbangkan penggunaan uji non-parametrik.")
                else:
                    narr = "Hasil uji Shapiro-Wilk menunjukkan bahwa seluruh variabel tidak berdistribusi normal (p < .05). Pertimbangkan penggunaan metode non-parametrik atau transformasi data."
            except Exception:
                narr = ""
        _add_ai_narasi(doc, profile, narr, "Interpretasi Normalitas")

    # BAB 4: Validitas
    if val_df is not None:
        try:
            val_empty = val_df.empty if hasattr(val_df,"empty") else not val_df
        except Exception:
            val_empty = True
        if not val_empty:
            _add_heading(doc, profile,
                         f"{bab_ctr}. Uji Validitas (Korelasi Pearson)", level=1)
            bab_no = bab_ctr; bab_ctr += 1
            _add_para(doc, profile,
                f"Butir dinyatakan valid apabila r-hitung \u2265 r-tabel = {r_tabel:.3f} (p < .05).",
                first_indent=True, is_last_in_block=True)
            df_val = val_df.to_frame() if isinstance(val_df, pd.Series) else pd.DataFrame(val_df)
            _style_table(doc, df_val, profile, f"Tabel {bab_no}.1. Hasil Uji Validitas Pearson")
            narr = ai_texts.get("validity","")
            if not narr and _has_narr:
                try: narr = narrate_validity(val_df, r_tabel).replace("**","")
                except Exception: narr = ""
            if not narr:
                try:
                    r_col  = next((c for c in df_val.columns if "r-hitung" in c.lower() or "r_hitung" in c.lower()), None)
                    ket_col= next((c for c in df_val.columns if "ket" in c.lower()), None)
                    if ket_col:
                        n_valid = (df_val[ket_col].str.lower() == "valid").sum()
                        n_total = len(df_val)
                        narr = (f"Seluruh {n_total} butir instrumen dinyatakan valid (r-hitung \u2265 {r_tabel:.3f})."
                                if n_valid == n_total else
                                f"Dari {n_total} butir instrumen, {n_valid} butir dinyatakan valid dan {n_total-n_valid} butir tidak valid (r-hitung < {r_tabel:.3f}).")
                except Exception:
                    narr = ""
            _add_ai_narasi(doc, profile, narr, "Interpretasi Validitas")

    # BAB 5: Reliabilitas
    if alpha_val is not None:
        try:
            alpha_f = float(alpha_val)
            _add_heading(doc, profile,
                         f"{bab_ctr}. Uji Reliabilitas (Cronbach's Alpha)", level=1)
            bab_no = bab_ctr; bab_ctr += 1
            status = "reliabel" if alpha_f >= 0.7 else "tidak reliabel"
            _add_para(doc, profile,
                f"Hasil pengujian menunjukkan Cronbach's Alpha (\u03b1) = {alpha_f:.3f}, "
                f"sehingga instrumen dinyatakan {status} (\u03b1 \u2265 .70).",
                first_indent=True, is_last_in_block=True)
            _style_table(doc,
                pd.DataFrame([{
                    "Cronbach's Alpha (\u03b1)": round(alpha_f,4),
                    "Keterangan": "Reliabel" if alpha_f >= 0.7 else "Tidak Reliabel",
                }]), profile, f"Tabel {bab_no}.1. Hasil Uji Reliabilitas")
        except Exception:
            pass

    # BAB 6: Korelasi
    if corr_matrix is not None:
        try:
            corr_empty = corr_matrix.empty
        except Exception:
            corr_empty = True
        if not corr_empty:
            _add_heading(doc, profile,
                         f"{bab_ctr}. Analisis Korelasi (Pearson)", level=1)
            bab_no = bab_ctr; bab_ctr += 1
            _add_para(doc, profile,
                "Matriks korelasi Pearson menggambarkan hubungan linier antar variabel.",
                first_indent=True, is_last_in_block=True)
            corr_display = corr_matrix.round(3).reset_index().rename(
                columns={"index":"Variabel"})
            _style_table(doc, corr_display, profile,
                         f"Tabel {bab_no}.1. Matriks Korelasi Pearson")
            narr = ai_texts.get("correlation","")
            if not narr:
                try:
                    pairs = []
                    cols_c = corr_matrix.columns.tolist()
                    for i, c1 in enumerate(cols_c):
                        for c2 in cols_c[i+1:]:
                            r = corr_matrix.loc[c1,c2]
                            if abs(r) >= 0.5:
                                arah = "positif" if r > 0 else "negatif"
                                pairs.append(f"{c1}\u2013{c2} (r = {r:.3f}, {arah})")
                    if pairs:
                        narr = ("Terdapat korelasi yang kuat (|r| \u2265 .50) antara: " +
                                "; ".join(pairs) + ". "
                                "Korelasi antar variabel prediktor perlu diperhatikan untuk menghindari masalah multikolinearitas.")
                    else:
                        narr = "Tidak ditemukan korelasi yang kuat (|r| \u2265 .50) antar variabel. Hubungan antar variabel tergolong lemah hingga sedang."
                except Exception:
                    narr = ""
            _add_ai_narasi(doc, profile, narr, "Interpretasi Korelasi")
            if figs_png.get("heatmap"):
                _embed_image(doc, profile, figs_png["heatmap"],
                             f"Gambar {fig_ctr}. Heatmap Korelasi Antar Variabel"); fig_ctr += 1
            sc = ai_texts.get("_scatter_meta", {})
            # Fix 6: accept both "scatter" and "korelasi_scatter" keys
            _scatter_png = figs_png.get("scatter") or figs_png.get("korelasi_scatter")
            if sc and _scatter_png:
                vx, vy = sc.get("var_x","X"), sc.get("var_y","Y")
                rv, pv = sc.get("r_val"), sc.get("p_val")
                _add_heading(doc, profile, f"{bab_no}.2. Scatter Plot: {vx} dan {vy}", level=2)
                if rv is not None and pv is not None:
                    sig = "signifikan" if float(pv) < 0.05 else "tidak signifikan"
                    _add_para(doc, profile,
                        f"Korelasi antara {vx} dan {vy} sebesar r = {float(rv):.3f} "
                        f"(p = {float(pv):.3f}), yang berarti {sig}.",
                        first_indent=True, is_last_in_block=True)
                _embed_image(doc, profile,
                             figs_png.get("scatter") or figs_png.get("korelasi_scatter"),
                             f"Gambar {fig_ctr}. Scatter Plot {vx} vs {vy}"); fig_ctr += 1

    # BAB 7+: Modul Lanjutan
    for mod_key, info in session_results.items():
        mod_data = info.get("data",{}) if isinstance(info,dict) else {}
        if not mod_data:
            continue
        renderer = _MODULE_RENDERERS.get(mod_key)
        if renderer is None:
            continue
        fig_ctr = renderer(doc, profile, mod_data, figs_png, fig_ctr, bab_ctr, ai_texts)
        bab_ctr += 1

    # Kesimpulan
    if ai_texts.get("kesimpulan"):
        _add_heading(doc, profile,
                     f"{bab_ctr}. Kesimpulan dan Rekomendasi", level=1)
        paras = [l.strip() for l in _clean_ai_text(ai_texts["kesimpulan"]).split("\n") if l.strip()]
        for i, pt in enumerate(paras):
            _add_para(doc, profile, pt, first_indent=True,
                      is_last_in_block=(i == len(paras)-1))

    # Referensi
    apa_refs = ai_texts.get("apa_references")
    if apa_refs:
        doc.add_page_break()
        _add_references_section(doc, profile, apa_refs)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# =============================================================================
# MARKDOWN GENERATOR
# =============================================================================

def generate_markdown_report(
    df, report, desc_df, norm_df, val_df, alpha_val, corr_matrix,
    cols, r_tabel,
    ai_texts=None, report_style="APA 7th Edition",
    data_type="Data Primer (Kuesioner / Survei)",
    session_results=None,
    user_name: str = "",
) -> str:
    if ai_texts        is None: ai_texts        = {}
    if session_results is None: session_results = {}
    lines = [
        "# Laporan Analisis Statistik","",
        "**Ruang Statistika \u2014 AI-Powered Research & Stats Reporting**","",
        "| Keterangan | Detail |","|---|---|",
        f"| Format | {report_style} |",
        f"| Tipe Data | {data_type.split('(')[0].strip()} |",
        f"| Total Responden | {report.get('rows_after_clean', len(df))} |",
        f"| Variabel Analisis | {len(cols)} |",
        *([ f"| Peneliti | {user_name} |" ] if user_name else []),
        f"| Tanggal | {datetime.date.today().strftime('%d %B %Y')} |",
        "","---","",
    ]
    bab = 1
    lines += [f"## {bab}. Ringkasan Data","",
              f"Dataset awal: **{report.get('original_rows','?')}** baris. "
              f"Setelah pembersihan: **{report.get('rows_after_clean', len(df))}** baris valid.",""]
    missing = report.get("missing_per_col",{})
    lines.append("Missing values: " + ", ".join([f"**{k}** ({v})" for k,v in missing.items()])
                 if missing else "Tidak ditemukan missing values.")
    bab += 1

    if desc_df is not None and not desc_df.empty:
        lines += ["", f"## {bab}. Statistik Deskriptif",""]
        lines.append(desc_df.to_markdown(index=False))
        narr = ai_texts.get("descriptive","")
        if narr:
            lines += ["","> **Interpretasi:**",""]
            for l in narr.split("\n"):
                if l.strip(): lines.append(f"> {l}")
        bab += 1

    if norm_df is not None and not norm_df.empty:
        lines += ["", f"## {bab}. Uji Normalitas","",
                  "H\u2080: data berdistribusi normal. Ditolak jika p < .05.",""]
        lines.append(norm_df.to_markdown(index=False))
        narr = ai_texts.get("normality","")
        if narr:
            lines += ["","> **Interpretasi:**",""]
            for l in narr.split("\n"):
                if l.strip(): lines.append(f"> {l}")
        bab += 1

    if val_df is not None:
        try:
            val_empty = val_df.empty if hasattr(val_df,"empty") else not val_df
        except Exception:
            val_empty = True
        if not val_empty:
            lines += ["", f"## {bab}. Uji Validitas","",
                      f"Butir valid: r-hitung \u2265 r-tabel = **{r_tabel}**",""]
            try: lines.append(val_df.to_markdown(index=False))
            except Exception: lines.append(str(val_df))
            bab += 1

    if alpha_val is not None:
        try:
            af = float(alpha_val)
            lines += ["", f"## {bab}. Uji Reliabilitas","",
                      f"**Cronbach's Alpha (\u03b1) = {af:.4f}** \u2014 "
                      f"[{'RELIABEL' if af >= 0.7 else 'TIDAK RELIABEL'}]",""]
            bab += 1
        except Exception: pass

    if corr_matrix is not None:
        try:
            corr_empty = corr_matrix.empty
        except Exception:
            corr_empty = True
        if not corr_empty:
            lines += ["", f"## {bab}. Analisis Korelasi",""]
            corr_md = corr_matrix.round(3).reset_index().rename(columns={"index":"Variabel"})
            lines.append(corr_md.to_markdown(index=False))
            narr = ai_texts.get("correlation","")
            if narr:
                lines += ["","> **Interpretasi:**",""]
                for l in narr.split("\n"):
                    if l.strip(): lines.append(f"> {l}")
            bab += 1

    _MD_LABELS = {
        "regresi":"Analisis Regresi Linier","ols_plus":"Regresi OLS (Diagnostik)",
        "logistik":"Regresi Logistik","mediasi":"Analisis Mediasi",
        "moderasi":"Analisis Moderasi","anova":"ANOVA & Post-hoc",
        "uji_beda":"Uji Beda","outlier":"Deteksi Outlier",
        "sem":"SEM & CFA","kelompok":"Analisis Kelompok",
    }
    for mod_key, info in session_results.items():
        mod_data = info.get("data",{}) if isinstance(info,dict) else {}
        if not mod_data: continue
        sec_title = _MD_LABELS.get(mod_key, info.get("label", mod_key))
        lines += ["", f"## {bab}. {sec_title}",""]
        for data_key, label in [
            ("coef_table","Koefisien"),("anova_table","Tabel ANOVA"),
            ("posthoc_table","Post-hoc"),("path_table","Koefisien Jalur"),
            ("fit_indices","Indeks Kecocokan"),("loadings","Factor Loadings"),
            ("path_estimates","Estimasi Jalur"),("outlier_table","Daftar Outlier"),
            ("group_stats","Statistik Kelompok"),
        ]:
            sub_df = mod_data.get(data_key)
            if sub_df is not None and not sub_df.empty:
                lines += [f"**{label}**",""]
                try: lines.append(sub_df.round(4).to_markdown(index=False))
                except Exception: lines.append(str(sub_df))
                lines.append("")
        # Fallback narasi MD
        narr = ai_texts.get(mod_key,"")
        if not narr:
            try: narr = _fallback_narasi(mod_key, mod_data)
            except Exception: narr = ""
        if narr:
            lines += ["",f"> **Interpretasi \u2014 {sec_title}:**",""]
            for l in _clean_ai_text(narr).split("\n"):
                if l.strip(): lines.append(f"> {l}")
            lines.append("")
        bab += 1

    if ai_texts.get("kesimpulan"):
        lines += ["","---",f"## {bab}. Kesimpulan dan Rekomendasi","",
                  _clean_ai_text(ai_texts["kesimpulan"])]

    apa_refs = ai_texts.get("apa_references")
    if apa_refs:
        lines += ["","---","## Referensi",""]
        refs = apa_refs if isinstance(apa_refs,list) else \
               [l for l in str(apa_refs).split("\n") if l.strip()]
        for r in refs: lines.append(f"{r.strip()}  ")

    lines += ["","---",
              "*Laporan dihasilkan otomatis oleh Ruang Statistika*",
              f"*Format: {report_style} | Powered by Python, Streamlit & AI*"]
    return "\n".join(lines)
